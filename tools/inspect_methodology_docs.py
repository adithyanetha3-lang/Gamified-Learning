from pathlib import Path
import json

from docx import Document


FILES = [
    Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/Methodology Version for Mandhamari.docx"),
    Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/Methodology Version for Mandhamari (1).docx"),
]


def inspect_doc(path):
    doc = Document(path)
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paras.append({"style": p.style.name, "text": text})
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(rows)
    return {
        "file": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "paragraphs": paras,
        "tables": tables,
    }


def main():
    out = [inspect_doc(path) for path in FILES]
    Path("outputs/pavan_mandamarri").mkdir(parents=True, exist_ok=True)
    Path("outputs/pavan_mandamarri/methodology_docs_extract.json").write_text(
        json.dumps(out, indent=2, ensure_ascii=False), encoding="utf-8"
    )
    for doc in out:
        print("\n===", Path(doc["file"]).name, "===")
        print("paragraphs", doc["paragraph_count"], "tables", doc["table_count"])
        for item in doc["paragraphs"][:80]:
            print(f"[{item['style']}] {item['text'][:220]}")
        print("TABLES")
        for table in doc["tables"][:3]:
            safe = json.dumps(table[:4], ensure_ascii=True)
            print(safe[:2000])


if __name__ == "__main__":
    main()
