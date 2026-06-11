from pathlib import Path
import json

import openpyxl
from docx import Document


XLSX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/PavanThesis.0.xlsx")
DOCX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/Duguta Sreeja Paper Draft.final version.docx")


def inspect_xlsx():
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    out = {"sheets": []}
    for ws in wb.worksheets:
        rows = list(ws.iter_rows(values_only=True))
        nonempty = [
            [cell for cell in row]
            for row in rows
            if any(cell is not None and str(cell).strip() for cell in row)
        ]
        headers = []
        if nonempty:
            for cell in nonempty[0]:
                headers.append("" if cell is None else str(cell).strip())
        out["sheets"].append(
            {
                "name": ws.title,
                "max_row": ws.max_row,
                "max_col": ws.max_column,
                "nonempty_rows": len(nonempty),
                "headers": headers[:25],
                "sample_rows": nonempty[1:6],
            }
        )
    return out


def inspect_docx():
    doc = Document(DOCX)
    paras = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paras.append({"style": p.style.name, "text": text[:220]})
        if len(paras) >= 80:
            break
    tables = []
    for table in doc.tables[:5]:
        table_rows = []
        for row in table.rows[:5]:
            table_rows.append([cell.text.strip()[:80] for cell in row.cells])
        tables.append(table_rows)
    return {"paragraphs": paras, "tables": tables, "paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables)}


def main():
    print(json.dumps({"xlsx": inspect_xlsx(), "docx": inspect_docx()}, indent=2, default=str))


if __name__ == "__main__":
    main()
