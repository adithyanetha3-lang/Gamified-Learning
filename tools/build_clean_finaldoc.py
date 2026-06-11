from pathlib import Path
from collections import Counter
import json
import re

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path("outputs/pavan_mandamarri")
OUT = BASE / "Finaldoc.docx"
ANALYSIS = BASE / "analysis.json"
XLSX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/PavanThesis.0.xlsx")
MAP = Path(r"C:/Users/adith/Downloads/ChatGPT Image Jun 3, 2026, 10_59_57 PM.png")
CHART_VULN = BASE / "vulnerability_levels.png"
CHART_EXP = BASE / "environmental_exposure.png"
CHART_LIV = BASE / "livelihood_sources.png"


def clean(v):
    if v is None:
        return ""
    return re.sub(r"\s+", " ", str(v).strip())


def rows_for(wb, sheet):
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    headers = [clean(x) or f"Column {i+1}" for i, x in enumerate(rows[0])]
    out = []
    for row in rows[1:]:
        if any(clean(x) for x in row):
            out.append({headers[i]: row[i] if i < len(row) else None for i in range(len(headers))})
    return headers, out


def freq(rows, key, limit=20):
    c = Counter(clean(row.get(key)) for row in rows if clean(row.get(key)))
    total = sum(c.values()) or 1
    return [[k, v, f"{v * 100 / total:.1f}%"] for k, v in c.most_common(limit)]


def stats(rows, key):
    vals = []
    for row in rows:
        v = row.get(key)
        if isinstance(v, (int, float)):
            vals.append(float(v))
        elif clean(v):
            m = re.search(r"\d+(?:,\d{3})*(?:\.\d+)?", clean(v))
            if m:
                vals.append(float(m.group(0).replace(",", "")))
    if not vals:
        return []
    vals = sorted(vals)
    n = len(vals)
    median = vals[n // 2] if n % 2 else (vals[n // 2 - 1] + vals[n // 2]) / 2
    return [["Count", n], ["Minimum", vals[0]], ["Maximum", vals[-1]], ["Mean", f"{sum(vals) / n:.2f}"], ["Median", f"{median:.2f}"]]


def yes_no_summary(rows, key):
    yes = no = other = 0
    for row in rows:
        v = clean(row.get(key)).lower()
        if v.startswith("yes"):
            yes += 1
        elif v.startswith("no"):
            no += 1
        elif v:
            other += 1
    total = yes + no + other or 1
    return yes, no, other, f"{yes * 100 / total:.1f}%"


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.4)
    sec.bottom_margin = Cm(2.1)
    sec.left_margin = Cm(2.6)
    sec.right_margin = Cm(2.1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for s in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[s].font.name = "Times New Roman"
        styles[s].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def shade(cell, color):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    pr.append(shd)


def cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(str(text))
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.28
    p.paragraph_format.space_after = Pt(7)
    p.add_run(text)
    return p


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell_text(cell, h, bold=True, color="FFFFFF")
        shade(cell, "1F4E79")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cell_text(cells[i], val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def picture(doc, path, caption, width=6.1):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True


def open_themes(wb):
    _, rows = rows_for(wb, "Open-ended")
    problems = " ".join(clean(r.get("J1. What do you think are the main problems your household faces because of mining?")).lower() for r in rows)
    support = " ".join(clean(r.get("J2. What support or change would help your household the most in reducing these problems?")).lower() for r in rows)
    problem_keys = {
        "Water scarcity and unsafe/unfiltered water": ["water scarcity", "unfiltered water", "water"],
        "Dust and air pollution": ["dust", "air pollution"],
        "Blasting vibration and cracks in houses": ["crack", "vibration", "blasting"],
        "Damaged roads and weak transport": ["road", "transport"],
        "Limited job sources and local employment": ["job", "employment"],
        "Health problems such as allergy/skin disease": ["allergy", "skin", "disease", "health"],
    }
    support_keys = {
        "Regular safe drinking-water supply/filter bed": ["water", "filter", "godavari"],
        "Local employment priority/reservation": ["job", "employment", "reservation"],
        "Compensation for house or infrastructure damage": ["compensation", "damage", "house"],
        "Road, transport, school and park facilities": ["road", "transport", "school", "park"],
        "Basic amenities and settlement improvements": ["amenities", "amenties", "land"],
    }
    p_rows = [[k, sum(problems.count(x) for x in keys)] for k, keys in problem_keys.items()]
    s_rows = [[k, sum(support.count(x) for x in keys)] for k, keys in support_keys.items()]
    return sorted(p_rows, key=lambda x: x[1], reverse=True), sorted(s_rows, key=lambda x: x[1], reverse=True), rows


def title_pages(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IMPACT OF MINING ON LOCAL COMMUNITIES\nMANDAMARRI, MANCHERIAL DISTRICT, TELANGANA")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Finaldoc - Ready to Submit Project Report")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph()
    table(doc, ["Particular", "Details"], [
        ["Submitted by", "Durgam Pavan"],
        ["Study Area", "Mandamarri, Mancherial District, Telangana"],
        ["Sample Size", "60 households"],
        ["Method", "Household survey, interviews, observations, vulnerability index and open-ended responses"],
        ["Date", "June 2026"],
    ], [4, 11])
    doc.add_page_break()
    heading(doc, "Declaration", 1)
    para(doc, "I hereby declare that this project report titled Impact of Mining on Local Communities: A Study of Mandamarri, Mancherial District is based on primary household survey data, field observations, methodological documents, and academic interpretation. The information, analysis, tables, recommendations, and conclusion are prepared for academic submission.")
    doc.add_page_break()
    heading(doc, "Acknowledgement", 1)
    para(doc, "I express my sincere gratitude to the respondents of Mandamarri who shared their time, experiences and concerns about mining, livelihood, health, water, transport and local development. Their participation made this study possible. I also thank my teachers, classmates and family members for their encouragement and guidance during the preparation of this project.")
    doc.add_page_break()


def abstract_toc(doc):
    heading(doc, "Abstract", 1)
    para(doc, "Mining is one of the most important economic activities in the coal belt of Mancherial district, Telangana. It provides employment, supports local markets and contributes to industrial development. At the same time, mining also creates environmental, social, health and livelihood-related pressures for the local communities that live near mining areas. This project examines the impact of mining on 60 households in Mandamarri, with special focus on livelihood dependence, environmental exposure, sensitivity, adaptive capacity and household-level vulnerability.")
    para(doc, "The study is based on primary survey data, personal interviews, field observations and open-ended responses. The methodology follows a mixed-method design and applies the IPCC vulnerability framework along with the Pressure and Release model. A Composite Vulnerability Index was used to combine Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index. The findings show that environmental exposure is very high: all surveyed households reported dust pollution, regular noise disturbance and water-quality problems. Around 66.7 percent of households are currently linked to mining work, but only 8.3 percent reported savings and no household reported receiving mining compensation.")
    para(doc, "The study concludes that mining in Mandamarri creates both opportunities and risks. Employment benefits exist, but they are accompanied by dust, noise, water-quality concerns, house cracks, road damage, health insecurity and weak compensation systems. The report recommends safe drinking-water supply, dust control, health camps, house-damage assessment, local employment priority, social protection, alternative livelihood support and participatory grievance redress.")
    para(doc, "Keywords: Mining impact, Mandamarri, Mancherial, coal mining, vulnerability, livelihood, exposure, adaptive capacity, local communities.")
    doc.add_page_break()
    heading(doc, "Table of Contents", 1)
    numbered(doc, [
        "Chapter 1: Introduction",
        "Chapter 2: Study Area Profile",
        "Chapter 3: Review of Literature and Conceptual Framework",
        "Chapter 4: Research Methodology",
        "Chapter 5: Results and Analysis",
        "Chapter 6: Discussion",
        "Chapter 7: Major Findings, Recommendations and Action Plan",
        "Chapter 8: Conclusion",
        "References",
        "Appendices",
    ])
    doc.add_page_break()


def chapter1(doc):
    heading(doc, "Chapter 1: Introduction", 1)
    heading(doc, "1.1 Background of the Study", 2)
    for text in [
        "Mining is a major economic activity in many parts of India and plays an important role in energy production, industrial growth and regional employment. Coal mining is especially significant in Telangana's northern coal belt, where mining towns have developed around extraction, transport, labour settlements and related services. Mandamarri in Mancherial district is one such mining-influenced town.",
        "Mining-led development is often viewed through production, revenue and employment. However, the experience of local communities is more complex. People living near mining areas may receive livelihood opportunities, but they may also face dust pollution, noise, blasting vibration, water-quality issues, damaged roads, health problems and uncertainty about compensation. These impacts enter daily household life through drinking water, housing, transport, health expenditure and employment security.",
        "The impact of mining on local communities therefore cannot be understood from one angle alone. It must be studied as a multidimensional process involving social, economic, environmental and institutional factors. A household may depend on mining wages and still suffer from mining-related pollution. A non-mining household may not receive direct employment benefits but may still bear environmental costs.",
        "Mandamarri provides an important setting for such a study because it is closely connected with coal mining and the Godavari Valley coalfield region. The town's economy, settlement pattern and livelihood structure are shaped by mining activity. At the same time, local households report concerns about water, dust, roads, blasting vibration, house cracks and lack of secure local employment.",
    ]:
        para(doc, text)
    heading(doc, "1.2 Problem Statement", 2)
    for text in [
        "The main problem examined in this study is the unequal and multidimensional impact of mining on local communities in Mandamarri. Mining provides employment and supports the local economy, but it also produces environmental and social pressures that affect quality of life. The community experiences mining not only as a source of work but also as a source of dust, noise, water stress, house damage and livelihood insecurity.",
        "In many mining settlements, the benefits of mining are visible in wages and infrastructure, while the costs are dispersed across households through health expenditure, repair costs, damaged roads, water scarcity and reduced environmental quality. This creates a development paradox: mining may support regional growth while increasing vulnerability among the people living closest to extraction sites.",
        "The problem is especially important because vulnerability is not equal among households. Families with illness, dependents, debt, unstable income, no land, no savings and weak social protection are more affected by the same mining environment than households with stronger assets and income security. Therefore, the study asks not only what mining impacts exist, but also which conditions make households more vulnerable.",
    ]:
        para(doc, text)
    heading(doc, "1.3 Aim of the Study", 2)
    para(doc, "The aim of the study is to assess the impact of mining on local communities in Mandamarri, Mancherial district, with particular focus on livelihood conditions, environmental exposure, health issues, socio-economic sensitivity, adaptive capacity and household vulnerability.")
    heading(doc, "1.4 Objectives of the Study", 2)
    bullets(doc, [
        "To examine the socio-economic profile of mining-affected households in Mandamarri.",
        "To analyse the role of mining in local livelihood, income and employment security.",
        "To identify major environmental impacts such as dust, noise, water-quality problems, road damage and blasting-related effects.",
        "To study household health conditions, debt, savings, dependents and coping capacity.",
        "To construct and interpret a Composite Vulnerability Index using exposure, sensitivity and adaptive capacity indicators.",
        "To understand community perceptions and support expected from mining authorities and government institutions.",
        "To suggest practical recommendations for reducing mining-related vulnerability and improving local well-being.",
    ])
    heading(doc, "1.5 Research Questions", 2)
    bullets(doc, [
        "What are the major social, economic and environmental impacts of mining on households in Mandamarri?",
        "How does mining influence livelihood dependence and income stability?",
        "What forms of environmental exposure are most commonly reported by local households?",
        "Which household conditions increase sensitivity to mining-related impacts?",
        "What resources and services strengthen or weaken adaptive capacity?",
        "What policy measures are needed to make mining development more community-centred?",
    ])
    heading(doc, "1.6 Significance of the Study", 2)
    for text in [
        "This study is significant because it focuses on the lived experience of households in a mining town. It moves beyond general statements about mining and uses primary data to identify specific issues faced by local communities. The findings can help students, researchers, local authorities and community organizations understand the relationship between mining, development and vulnerability.",
        "The study is also useful because it combines quantitative and qualitative evidence. Frequencies and index scores show the scale of problems, while open-ended responses explain how households experience those problems. This makes the project both data-based and community-sensitive.",
        "The recommendations are practical and locally relevant. They focus on safe water, dust control, health protection, compensation, road repair, local employment, grievance redress and livelihood diversification. These are issues that directly emerge from the survey responses.",
    ]:
        para(doc, text)
    heading(doc, "1.7 Scope of the Study", 2)
    para(doc, "The scope of the study is limited to 60 households surveyed in Mandamarri. It examines household-level impact and vulnerability rather than technical mine production data. The project does not claim to represent all mining regions, but it provides a detailed local understanding of mining impact in one important coal mining settlement.")
    heading(doc, "1.8 Organization of the Report", 2)
    para(doc, "The report is organized into eight chapters. Chapter 1 introduces the research problem, objectives and significance. Chapter 2 presents the study area profile. Chapter 3 reviews relevant literature and conceptual frameworks. Chapter 4 explains the methodology. Chapter 5 presents results and analysis. Chapter 6 discusses the findings in depth. Chapter 7 provides major findings, recommendations and an action plan. Chapter 8 concludes the project.")
    doc.add_page_break()


def chapter2(doc):
    heading(doc, "Chapter 2: Study Area Profile", 1)
    heading(doc, "2.1 Location and Administrative Profile", 2)
    for text in [
        "Mandamarri is a prominent mining town and mandal located in Mancherial district of Telangana. It is part of the northern coal belt region of the state and is widely known for its association with coal mining. The settlement is linked with nearby towns such as Mancherial, Bellampalli, Chennur, Luxettipet and Dandepalli, forming part of a wider mining-influenced regional economy.",
        "Administratively, Mandamarri's position in Mancherial district is important because the district itself has developed as a significant coal mining and industrial region. The town's location near road and railway networks supports movement of coal, workers, goods and services. These connections strengthen the mining economy but also bring traffic, dust and road-use pressure.",
    ]:
        para(doc, text)
    picture(doc, MAP, "Figure 2.1: Location map of Mandamarri, Mancherial district, Telangana, India.", 6.3)
    heading(doc, "2.2 Physical Geography and Coalfield Setting", 2)
    for text in [
        "Mandamarri is associated with the Godavari Valley coalfield region, one of India's important coal-bearing zones. The physical landscape includes undulating terrain, dry deciduous vegetation, seasonal climatic conditions and landforms influenced by both natural processes and industrial activity.",
        "The geological characteristics of the region have played a decisive role in its economic development. Coal-bearing formations made the area suitable for extraction, which in turn shaped settlement growth, labour migration and infrastructure development. Mining has altered land use, vegetation patterns, drainage conditions and the physical appearance of the landscape.",
        "The proximity of Mandamarri to river systems and transport corridors makes environmental management essential. Dust, water movement, drainage disturbance, road damage and settlement expansion are not isolated technical problems; they directly affect households through health, water access, mobility and housing conditions.",
    ]:
        para(doc, text)
    heading(doc, "2.3 Historical Development of Mining Activities", 2)
    for text in [
        "The development of Mandamarri is closely connected with the history of coal mining in the Godavari Valley coalfields. As coal extraction expanded, mining infrastructure, labour colonies, transport routes and market services developed around the mines. The settlement gradually became economically dependent on mining and related activities.",
        "Mining attracted workers from different socio-economic and cultural backgrounds. Over time, Mandamarri developed a working-class social structure that includes permanent mine workers, contract workers, informal labourers, service workers, vendors, drivers and small business households. This history explains why mining remains central to local identity.",
        "In recent decades, mining activity has continued through underground and opencast operations and modernization of extraction methods. These changes have increased production capacity but also intensified concerns about dust, heavy vehicles, land disturbance, blasting vibration and local employment security.",
    ]:
        para(doc, text)
    heading(doc, "2.4 Demographic and Social Characteristics", 2)
    for text in [
        "The demographic profile of Mandamarri reflects its mining history and settlement development. The local population includes Scheduled Castes, Scheduled Tribes, Backward Classes, other caste groups and religious minorities. Migration and long-term residence associated with mining have contributed to social diversity.",
        "The household survey shows that respondents belong to different caste and livelihood categories. This diversity is important because mining impacts are filtered through social and economic conditions. A household with stable income, savings and better health may cope differently from a household with debt, illness and dependents.",
        "The presence of dependents in many households increases the need for stable income, health protection, education facilities and safe living conditions. Mining impacts therefore affect not only workers but entire families.",
    ]:
        para(doc, text)
    heading(doc, "2.5 Economic Structure and Livelihood Pattern", 2)
    for text in [
        "The economy of Mandamarri is predominantly shaped by coal mining and related activities. Mining provides direct employment to workers and indirect income to transport operators, vendors, shops, mechanics, daily wage labourers and service workers. Singareni-linked mining activity is a major economic influence in the region.",
        "Despite the dominance of mining, other livelihoods continue to exist. Agriculture, small businesses, driving, private jobs, security work, vending and daily wage labour form part of the local economy. This mixed livelihood structure shows that mining influences both direct employees and non-mining households.",
        "The survey indicates that mining-linked work is important, but livelihood security is uneven. Some households have stable employment, while others depend on contract work, seasonal activity or one earning member. This unevenness is central to understanding vulnerability.",
    ]:
        para(doc, text)
    heading(doc, "2.6 Environmental Conditions", 2)
    for text in [
        "The environmental conditions of Mandamarri have been shaped by prolonged mining activity. Dust generated from blasting, excavation, coal transport and heavy vehicle movement is one of the most visible impacts. Noise from machinery, vehicles and blasting also affects the settlement environment.",
        "Water resources are a major concern in the study area. Respondents reported water-quality problems, and open-ended responses repeatedly mentioned unfiltered water, water scarcity and the need for safe drinking-water supply. This shows that water is both an environmental and public health issue.",
        "Mining-related environmental pressure also appears through road damage, loss of vegetation, house cracks and anxiety about blasting vibration. These impacts demonstrate that environmental change becomes a household welfare issue when it affects health, mobility, housing and daily living conditions.",
    ]:
        para(doc, text)
    doc.add_page_break()


def chapter3(doc):
    heading(doc, "Chapter 3: Review of Literature and Conceptual Framework", 1)
    sections = [
        ("3.1 Mining and the Development Paradox", [
            "Literature on mining frequently discusses a development paradox. Mining contributes to economic growth, employment, energy production and industrial expansion, but local communities living near extraction sites often experience environmental degradation, displacement pressures, health risks and social insecurity.",
            "This paradox is highly relevant to Mandamarri. Mining supports local livelihoods, yet households report dust, noise, water-quality concerns, road damage and house cracks. A development activity therefore produces both benefits and burdens.",
        ]),
        ("3.2 Vulnerability as a Concept", [
            "Vulnerability has emerged as an important concept for understanding how communities encounter risk and transformation. Earlier studies often focused on hazards as external events, but later social research shows that risk is shaped by poverty, assets, power, institutions and coping capacity.",
            "Vulnerability does not mean exposure alone. A household becomes vulnerable when exposure combines with sensitivity and weak adaptive capacity. For example, dust pollution may affect all households, but a household with illness, debt and no savings may suffer more severely.",
        ]),
        ("3.3 IPCC Vulnerability Framework", [
            "The IPCC framework explains vulnerability through exposure, sensitivity and adaptive capacity. Exposure refers to contact with hazards or stressors. Sensitivity refers to the degree to which a system or household is affected. Adaptive capacity refers to the ability to adjust, cope and recover.",
            "In this project, exposure includes dust, noise, water-quality problems, accident risk and proximity to mines. Sensitivity includes debt, dependents, illness, landlessness and unstable income. Adaptive capacity includes savings, bank accounts, group membership, social protection, compensation, health facilities and schools.",
        ]),
        ("3.4 Pressure and Release Model", [
            "The Pressure and Release model explains vulnerability through root causes, dynamic pressures and unsafe conditions. Root causes are deeper structural conditions such as dependence on mining, inequality and weak institutional support. Dynamic pressures translate these into everyday risk through poor services, contract work, weak enforcement and limited alternatives.",
            "Unsafe conditions are the immediate circumstances in which people live, such as dust exposure, weak housing, damaged roads, unsafe water and health risks. This model is useful because it connects household-level findings with broader social processes.",
        ]),
        ("3.5 Vulnerability in Mining Communities", [
            "Studies of mining communities show that mining produces complex vulnerability. It can generate employment and infrastructure, but it can also create land degradation, pollution, water stress, occupational health risks, informal labour dependence and uneven benefit distribution.",
            "In mining settlements, vulnerability is not restricted to mine workers. Families, farmers, vendors, drivers, women, children and elderly people can all experience mining-related impacts. Therefore, local community studies are necessary for understanding the full social cost of mining.",
        ]),
        ("3.6 Conceptual Framework of the Present Study", [
            "The present study combines the IPCC framework and the PAR model. The IPCC framework helps measure exposure, sensitivity and adaptive capacity. The PAR model helps interpret root causes, dynamic pressures and unsafe conditions.",
            "This combined framework is effective for Mandamarri because the survey data show both measurable indicators and lived experiences. The Composite Vulnerability Index provides numerical classification, while open-ended responses and discussion explain the meaning of vulnerability.",
        ]),
    ]
    for h, paras in sections:
        heading(doc, h, 2)
        for text in paras:
            para(doc, text)
    table(doc, ["Framework", "Main idea", "Use in this study"], [
        ["IPCC vulnerability framework", "Vulnerability is shaped by exposure, sensitivity and adaptive capacity", "Used to construct EI, SI, ADI and CVI"],
        ["Pressure and Release model", "Risk emerges from root causes, dynamic pressures and unsafe conditions", "Used to interpret drivers of mining vulnerability"],
        ["Livelihood approach", "Household well-being depends on assets, income, capabilities and coping strategies", "Used to analyse mining work, agriculture, small business and daily wage dependence"],
        ["Environmental justice perspective", "Focuses on distribution of environmental burdens and benefits", "Used to discuss pollution, compensation and participation"],
    ], [4, 6, 6])
    doc.add_page_break()


def chapter4(doc):
    heading(doc, "Chapter 4: Research Methodology", 1)
    sections = [
        ("4.1 Methodological Orientation", [
            "The study follows a mixed-method approach because mining vulnerability includes both measurable and experiential dimensions. Quantitative data provide frequencies, percentages and vulnerability scores, while qualitative responses explain how households experience mining-related stress.",
            "The methodology is guided by the understanding that vulnerability is context-specific. The same mining environment can affect households differently depending on income, health, savings, dependents, employment type, debt and access to support systems.",
        ]),
        ("4.2 Research Design", [
            "The research design is descriptive and analytical. It is descriptive because it records the socio-economic, health, livelihood and environmental profile of households. It is analytical because it uses indicators and index construction to interpret household vulnerability.",
            "The design is suitable for a local project because it captures both the structure of household conditions and the meaning of mining impacts. It allows the study to move from individual responses to wider conclusions about community vulnerability.",
        ]),
        ("4.3 Sources of Data", [
            "Primary data were collected from 60 households in Mandamarri through a structured household survey. The questionnaire covered basic information, education, health, livelihood, exposure, sensitivity, adaptive capacity, shocks, perceptions, assets and open-ended responses.",
            "Personal interviews and field observations supported the survey. These methods helped clarify responses and capture visible conditions such as housing, roads, water concerns and environmental pressure. Secondary understanding was drawn from academic frameworks, mining-community literature and vulnerability methodology.",
        ]),
        ("4.4 Sampling Procedure", [
            "Purposive sampling was adopted because the study focuses specifically on mining-affected households in Mandamarri. The sample was selected to include households with different livelihood categories and social backgrounds.",
            "A total of 60 households were surveyed. Although purposive sampling does not allow statistical generalization to all mining regions, it is appropriate for a focused local vulnerability study because it captures households that are directly relevant to the research problem.",
        ]),
        ("4.5 Data Collection Methods", [
            "The main data collection tool was a semi-structured questionnaire. The structured part allowed quantitative analysis, while open-ended questions allowed respondents to describe problems and expectations in their own words.",
            "Personal interviews were conducted during the survey process. These interviews helped identify issues such as unfiltered water, house cracks, road damage, blasting vibration, lack of transport and local job expectations. Field observation added contextual understanding to the numerical data.",
        ]),
        ("4.6 Indicator Selection", [
            "Indicators were selected based on theoretical relevance, data availability, variability and interpretability. Each indicator was placed under Exposure Index, Sensitivity Index or Adaptive Capacity Deficit Index.",
            "Exposure indicators include distance from mine, dust, noise, accident, water quality and agriculture/livestock impact. Sensitivity indicators include education, illness, income stability, debt, land, dependents and marginalization. Adaptive capacity indicators include savings, bank account, group membership, social protection, compensation, alternative livelihood, health facility and school access.",
        ]),
        ("4.7 Normalisation and Index Construction", [
            "Normalisation was required because indicators were measured in different forms such as binary responses, ordinal categories and numeric values. The min-max logic was used to bring indicators onto a common scale and make them comparable.",
            "Positive indicators increase vulnerability when their value rises, while negative indicators are reversed so that higher normalized values consistently indicate higher vulnerability. The Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index were then combined to create the Composite Vulnerability Index.",
        ]),
    ]
    for h, paras in sections:
        heading(doc, h, 2)
        for text in paras:
            para(doc, text)
    heading(doc, "4.8 Twelve-Step Vulnerability Assessment Protocol", 2)
    table(doc, ["Step", "Application in this study", "Purpose"], [
        ["1", "Defined research aim, Mandamarri study area and household-level vulnerability", "Clarifies scope"],
        ["2", "Selected mixed-method household assessment", "Combines survey and lived experience"],
        ["3", "Adopted Tier-2 composite-index method", "Creates measurable vulnerability scores"],
        ["4", "Defined sector as mining-affected local community", "Links household risk with mining context"],
        ["5", "Selected EI, SI and ADI indicators", "Covers multidimensional vulnerability"],
        ["6", "Coded responses from 60 households", "Prepares data for scoring"],
        ["7", "Normalized indicators", "Makes different indicators comparable"],
        ["8", "Applied weighting and sub-index grouping", "Builds index structure"],
        ["9", "Computed EI, SI, ADI and CVI", "Measures household vulnerability"],
        ["10", "Prepared tables, charts and map", "Improves presentation"],
        ["11", "Classified vulnerability levels", "Identifies low, medium and high risk households"],
        ["12", "Interpreted drivers through IPCC and PAR models", "Connects scores with social causes"],
    ], [1.5, 8, 6])
    heading(doc, "4.9 Limitations of the Study", 2)
    bullets(doc, [
        "The study is based on 60 households and is therefore a local-level assessment.",
        "The findings are based on self-reported survey responses and field interpretation.",
        "Technical environmental testing of air, water and vibration was not conducted.",
        "The study focuses on household vulnerability and does not measure mine production or company-level economics.",
    ])
    doc.add_page_break()


def chapter5(doc, data, wb):
    heading(doc, "Chapter 5: Results and Analysis", 1)
    _, basic = rows_for(wb, "Basic Information")
    _, edu = rows_for(wb, "Section_B_Education_Health")
    _, liv = rows_for(wb, "Section_C_Livelihood")
    _, exp = rows_for(wb, "Section_D_Exposure")
    _, sens = rows_for(wb, " Sensitivity (socio-economic v")
    _, adap = rows_for(wb, "Adaptive Capacity")
    _, shock = rows_for(wb, "Shocks, losses & coping strateg")
    _, percep = rows_for(wb, "Perceptions, awareness & partic")
    _, asset = rows_for(wb, "Assets & infrastructure")
    _, sheet10 = rows_for(wb, "Sheet10")
    problem_themes, support_themes, open_rows = open_themes(wb)

    heading(doc, "5.1 Demographic Profile of Respondents", 2)
    para(doc, "The survey covered 60 households in Mandamarri. The demographic profile shows a working-age population with family responsibilities. The mean age of respondents is 41.65 years and the average household size is 3.4 members. Male respondents form 90 percent of the sample and female respondents form 10 percent, reflecting the field pattern of household representation.")
    para(doc, "The caste profile includes Scheduled Castes, Scheduled Tribes, Backward Classes, Other Backward Classes and other caste groups. This diversity is important because social position influences access to resources, employment and institutional support.")
    table(doc, ["Statistic", "Age", "Household Size"], [[a, b, c] for (a, b), (_, c) in zip(stats(basic, "Age"), stats(basic, "HH Size"))], [5, 4, 4])
    table(doc, ["Gender", "Households", "Percentage"], freq(basic, "Gender"), [6, 4, 4])
    table(doc, ["Caste", "Households", "Percentage"], freq(basic, "Caste"), [6, 4, 4])
    table(doc, ["Housing Type", "Households", "Percentage"], freq(basic, "HH_type"), [6, 4, 4])

    heading(doc, "5.2 Education and Health Conditions", 2)
    para(doc, "Education is an important human-capital indicator. In the survey, 43.3 percent of respondents reported higher secondary education and 23.3 percent reported graduation. The presence of education and technical training creates potential for livelihood diversification, but this potential requires local employment opportunities and skill-based placement.")
    para(doc, "Health results reveal vulnerability. Illness affected earning capacity in 38.3 percent of households, hospitalization was reported by 35 percent and health insurance was reported by only 35 percent. This indicates that health shocks can reduce income while increasing household expenditure.")
    table(doc, ["Education Level", "Households", "Percentage"], freq(edu, "B1_Education"), [6, 4, 4])
    table(doc, ["Health Indicator", "Finding", "Interpretation"], [
        ["Vocational/technical training", "40 households (66.7%)", "Skill base exists but requires better employment linkage"],
        ["Illness affecting earning", "23 households (38.3%)", "Health directly affects income security"],
        ["Hospitalization", "21 households (35.0%)", "Medical shocks are significant"],
        ["Medical expense burden", "20 households (33.3%)", "Health costs increase sensitivity"],
        ["Health insurance", "21 households (35.0%)", "Insurance coverage is limited"],
    ], [5, 5, 6])

    heading(doc, "5.3 Livelihood and Income Pattern", 2)
    para(doc, "Livelihood data show strong dependence on mining and related activities. Around 66.7 percent of households reported current connection with mining work. The occupational structure includes mining workers, contract workers, coal cutters, trolley drivers, security guards, farmers, vendors, private job workers and daily wage labourers.")
    para(doc, "The average number of economically active members is only 1.1 per household. This means that many households depend on one main earner. If that earner faces illness, accident, job loss or wage interruption, the entire household becomes economically vulnerable.")
    picture(doc, CHART_LIV, "Figure 5.1: Main income sources among surveyed households.")
    table(doc, ["Main Livelihood", "Households", "Percentage"], freq(liv, "C1_Main_Livelihood", 18), [7, 4, 4])
    table(doc, ["Income Stability", "Households", "Percentage"], freq(liv, "C7_Income Stability"), [6, 4, 4])
    table(doc, ["Livelihood Indicator", "Result", "Meaning"], [
        ["Currently connected to mining", "40 households (66.7%)", "Mining is economically central"],
        ["Past mining connection", "3 households (5.0%)", "Some households have previous mining dependence"],
        ["Average active members", "1.1 per household", "High dependence on one earner"],
        ["Formal documentation", "46 households (76.7%)", "Administrative access exists for many households"],
        ["Seasonal/not stable income", "24 households (40.0%)", "Income vulnerability remains significant"],
    ], [5, 5, 6])

    heading(doc, "5.4 Environmental Exposure", 2)
    para(doc, "Environmental exposure is the strongest and most consistent result of the study. All surveyed households reported dust or airborne pollution, regular noise disturbance and water-quality problems. This indicates that environmental stress is not limited to a few households but forms a common settlement-level experience.")
    para(doc, "Open-ended responses support the quantitative data. Respondents mentioned unfiltered water, water scarcity in summer, damaged roads due to heavy vehicles, dust from opencast activity, blasting vibration and cracks in houses. These responses show that environmental impact is experienced through daily household problems.")
    picture(doc, CHART_EXP, "Figure 5.2: Environmental exposure reported by surveyed households.")
    table(doc, ["Exposure Indicator", "Households Reporting Impact", "Percentage"], [
        ["Dust/airborne pollution", "60", "100.0%"],
        ["Regular noise disturbance", "60", "100.0%"],
        ["Water-quality problem", "60", "100.0%"],
        ["Accident experience", "8", "13.3%"],
        ["Agriculture/livestock affected", "5", "8.3%"],
        ["Displacement/relocation", "0", "0.0%"],
    ], [6, 5, 4])
    table(doc, ["Distance from Mine", "Households", "Percentage"], freq(exp, "D1_Distance_Mine"), [6, 4, 4])

    heading(doc, "5.5 Sensitivity Conditions", 2)
    para(doc, "Sensitivity refers to household conditions that make the effect of exposure stronger. In Mandamarri, important sensitivity factors include landlessness, dependents, debt, illness, low savings and unstable income.")
    para(doc, "Sixty percent of households reported no land, 35 percent reported debt and 85 percent reported dependents. These conditions increase vulnerability because households have fewer assets, more financial pressure and greater responsibility.")
    table(doc, ["Sensitivity Indicator", "Finding", "Interpretation"], [
        ["No land", "36 households (60.0%)", "Weak asset base"],
        ["Dependents present", "51 households (85.0%)", "Higher family responsibility"],
        ["Debt", "21 households (35.0%)", "Financial pressure"],
        ["Livestock ownership", "5 households (8.3%)", "Limited supplementary livelihood asset"],
        ["Tap water in house", "53 households (88.3%)", "Access exists but quality concerns remain"],
        ["Private toilet", "60 households (100.0%)", "Sanitation facility is strong"],
    ], [5, 5, 6])
    table(doc, ["Loan Source", "Households", "Percentage"], freq(sens, "E5b_ Main source of loan"), [6, 4, 4])

    heading(doc, "5.6 Adaptive Capacity", 2)
    para(doc, "Adaptive capacity refers to the ability of households to cope with and recover from stress. The data show both strengths and weaknesses. All households reported bank accounts and access to a PHC and primary school. Group membership is also high, which indicates some social capital.")
    para(doc, "However, only 8.3 percent reported savings and no household reported receiving mining compensation. This is a major concern because savings and compensation are key buffers during illness, house damage or livelihood shock.")
    table(doc, ["Adaptive Capacity Indicator", "Finding", "Interpretation"], [
        ["Savings", "5 households (8.3%)", "Weak household financial buffer"],
        ["Bank account", "60 households (100.0%)", "Institutional access exists"],
        ["Group membership", "46 households (76.7%)", "Social capital is present"],
        ["Ration card/social protection", "40 households (66.7%)", "Basic welfare support exists"],
        ["Mining compensation", "0 households (0.0%)", "Major accountability gap"],
        ["PHC access", "60 households (100.0%)", "Health facility access exists"],
        ["Primary school access", "60 households (100.0%)", "Basic education access exists"],
    ], [5, 5, 6])

    heading(doc, "5.7 Shocks, Assets and Perceptions", 2)
    para(doc, "Shocks and coping strategies show how households respond when crisis occurs. Serious illness and economic shocks can push households toward borrowing or asset sale. Since savings are low, households may rely on banks, relatives or moneylenders during crisis.")
    para(doc, "Perception responses show that residents consider mining impacts severe. Many respondents demanded safe water, filter beds, local employment, compensation for house damage, road repair, transport improvement and better amenities.")
    table(doc, ["Most Recent Major Shock", "Households", "Percentage"], freq(shock, "G2_Most recent major shock"), [6, 4, 4])
    table(doc, ["Impact Severity", "Households", "Percentage"], freq(percep, "H1_Impact_Severity"), [6, 4, 4])
    table(doc, ["Transport Owned", "Households", "Percentage"], freq(asset, "I4. Means of transport owned"), [6, 4, 4])
    table(doc, ["Major Problem Theme", "Relative Mention Count"], problem_themes, [10, 4])
    table(doc, ["Support Expected by Respondents", "Relative Mention Count"], support_themes, [10, 4])

    heading(doc, "5.8 Composite Vulnerability Index", 2)
    para(doc, "The Composite Vulnerability Index combines the Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index. The overall vulnerability score ranges from 0.05 to 0.33, with an average of about 0.19. This shows moderate average vulnerability, with variation across households.")
    para(doc, "The index confirms that vulnerability is not produced by one factor alone. It emerges when environmental exposure combines with illness, debt, dependents, unstable income, weak savings and lack of compensation.")
    picture(doc, CHART_VULN, "Figure 5.3: Composite vulnerability levels among surveyed households.")
    table(doc, ["Index", "Mean", "Minimum", "Maximum", "Interpretation"], [
        ["Exposure Index", "0.17", "0.10", "0.28", "Environmental pressure is widespread"],
        ["Sensitivity Index", "0.13", "0.03", "0.30", "Household vulnerability varies"],
        ["Adaptive Capacity Deficit", "0.11", "0.05", "0.20", "Low savings and no compensation weaken resilience"],
        ["Overall Vulnerability Index", "0.19", "0.05", "0.33", "Moderate average vulnerability"],
    ], [4, 2.2, 2.2, 2.2, 6])
    table(doc, ["Vulnerability Level", "Households", "Percentage"], freq(sheet10, "Vulnerability Level"), [6, 4, 4])
    doc.add_page_break()


def chapter6(doc):
    heading(doc, "Chapter 6: Discussion", 1)
    sections = [
        ("6.1 Mining as Opportunity and Risk", [
            "The findings show that mining is both an opportunity and a risk for Mandamarri. It provides employment, income and local economic activity, but it also produces environmental stress and social vulnerability. This dual character is the central theme of the study.",
            "Mining-linked households may receive income, but they are not free from dust, noise, water problems and health concerns. Non-mining households may not receive secure employment benefits, yet they still share the environmental burden. This uneven distribution of benefits and costs is a major issue in mining-area development.",
            "The survey shows that employment alone cannot be treated as complete development. A household may have mining income but still lack savings, health security, compensation and safe living conditions. Therefore, development must be measured through well-being, not only jobs.",
        ]),
        ("6.2 Environmental Exposure as a Settlement-Level Problem", [
            "The most striking result is that dust, noise and water-quality problems were reported by all surveyed households. This suggests that environmental exposure is not a private problem of selected families but a common settlement-level condition.",
            "Dust affects the everyday environment by entering houses, roads, water sources and public spaces. Noise and blasting vibration disturb the sense of safety. Water-quality concerns affect health, hygiene and household management. Together, these issues reduce quality of life.",
            "Environmental management must therefore be treated as a core part of mining governance. Dust control, road cleaning, water sprinkling, covered coal transport, water testing and public reporting should be regular practices.",
        ]),
        ("6.3 Livelihood Dependence and Economic Vulnerability", [
            "The livelihood results show strong dependence on mining. However, dependence on a single industry can increase vulnerability when employment is unstable or when households have only one earning member. The average of 1.1 economically active members per household indicates this risk.",
            "Contract and informal work can be especially insecure. Even where income appears stable, households without savings may struggle during illness, accident or job interruption. Livelihood diversification is therefore necessary.",
            "Local skill training, youth employment, self-help group enterprises, small business support and non-mining livelihood options can reduce dependence and improve resilience.",
        ]),
        ("6.4 Health, Debt and Coping Capacity", [
            "Health vulnerability is closely connected with economic vulnerability. Illness affecting earning capacity, hospitalization and medical expenses can push households into debt. Since savings are low, households may borrow during crisis.",
            "Debt is not only a financial indicator; it is a sign of weak coping capacity. When households borrow for health, house repair or livelihood needs, repayment can reduce future consumption and increase stress.",
            "Regular health camps, insurance enrolment, occupational health monitoring and PHC strengthening are important for reducing household vulnerability in Mandamarri.",
        ]),
        ("6.5 Governance, Compensation and Environmental Justice", [
            "The absence of reported mining compensation is one of the most important governance findings. Respondents mention house cracks, road damage, water problems and environmental stress, but no household reported receiving mining compensation.",
            "Environmental justice requires fair distribution of benefits and burdens. If local communities carry the costs of mining without adequate redress, mining-led development becomes unequal. Compensation, grievance redress and participation are therefore essential.",
            "A transparent complaint system should record house damage, water problems, health concerns and road issues. Verification should be timely, and action taken should be communicated publicly.",
        ]),
        ("6.6 Vulnerability Drivers through IPCC and PAR Models", [
            "Using the IPCC framework, Mandamarri's vulnerability can be understood through high exposure, varied sensitivity and weak adaptive capacity. Exposure is high because dust, noise and water problems are widespread. Sensitivity differs due to debt, dependents, illness, landlessness and unstable income. Adaptive capacity is weakened by low savings and lack of compensation.",
            "Using the PAR model, root causes include mining dependence, limited alternative livelihoods and weak compensation systems. Dynamic pressures include contract work, health expenditure, debt and limited institutional response. Unsafe conditions include dust, water-quality problems, road damage, house cracks and noise.",
            "These frameworks show that vulnerability is not accidental. It is produced by the interaction of mining activity, household conditions and institutional response.",
        ]),
    ]
    for h, paras in sections:
        heading(doc, h, 2)
        for text in paras:
            para(doc, text)
    doc.add_page_break()


def chapter7(doc):
    heading(doc, "Chapter 7: Major Findings, Recommendations and Action Plan", 1)
    heading(doc, "7.1 Major Findings", 2)
    bullets(doc, [
        "Mining is central to the Mandamarri economy, with 66.7 percent of surveyed households currently linked to mining work.",
        "Environmental exposure is widespread: 100 percent of households reported dust pollution, regular noise disturbance and water-quality problems.",
        "The average number of economically active members is only 1.1 per household, indicating dependence on one main earner.",
        "Health insecurity is significant, with illness affecting earning capacity in 38.3 percent of households and hospitalization reported by 35 percent.",
        "Sensitivity is increased by landlessness, debt, dependents, illness and unstable income.",
        "Adaptive capacity is mixed: bank access, PHC access and school access exist, but savings are very low and mining compensation is absent.",
        "Open-ended responses identify water scarcity, unfiltered water, dust, blasting vibration, house cracks, damaged roads and lack of local employment as major problems.",
        "The Composite Vulnerability Index shows moderate average vulnerability, with higher risk among households facing overlapping environmental and socio-economic pressures.",
    ])
    heading(doc, "7.2 Recommendations", 2)
    bullets(doc, [
        "Provide safe drinking-water supply, filtration and regular water-quality testing in Mandamarri.",
        "Strengthen dust-control measures through water sprinkling, covered transport, road cleaning and green buffers.",
        "Conduct technical inspection of houses reporting cracks and blasting-related damage.",
        "Create a transparent compensation mechanism for verified mining-related damage.",
        "Organize regular health camps focusing on respiratory illness, skin disease, hearing, back pain and occupational risks.",
        "Increase health insurance coverage and support medical referral for vulnerable households.",
        "Give priority to local youth in mining-related jobs, contract work, apprenticeships and skill training.",
        "Promote alternative livelihoods through SHGs, small enterprises, driving, repair services, tailoring, livestock and service-sector training.",
        "Repair roads damaged by heavy mining vehicles and improve regular transport connectivity.",
        "Create a community grievance redress committee with residents, Gram Panchayat, mine authorities and district officials.",
    ])
    heading(doc, "7.3 Action Plan", 2)
    table(doc, ["Problem", "Evidence", "Action Required", "Responsible Actors"], [
        ["Water quality and scarcity", "100% reported water-quality problems", "Water testing, filtration, regular supply", "Gram Panchayat, district administration, mine authorities"],
        ["Dust pollution", "100% reported dust pollution", "Sprinkling, covered transport, road cleaning, monitoring", "Mine authorities, pollution control authorities"],
        ["Noise and blasting vibration", "100% reported regular noise; responses mention cracks", "Vibration monitoring, blasting schedule information", "Mine authorities, technical departments"],
        ["House damage", "Open-ended responses mention cracks", "Technical inspection and compensation", "Revenue department, mine authorities"],
        ["Health insecurity", "Illness and hospitalization reported", "Health camps and insurance enrolment", "Health department, PHC"],
        ["Low savings and debt", "Only 8.3% reported savings; 35% debt", "Financial literacy, SHG credit, low-interest support", "Banks, SHGs, welfare departments"],
        ["Local employment demand", "Respondents requested local job priority", "Skill training and local recruitment linkage", "SCCL, labour department, skill mission"],
        ["Weak grievance redress", "No compensation reported", "Formal complaint register and public review", "District administration, Gram Panchayat"],
    ], [3.5, 4.5, 5, 4])
    heading(doc, "7.4 Monitoring Indicators", 2)
    table(doc, ["Indicator", "Frequency", "Purpose"], [
        ["Water-quality test results", "Quarterly", "Track drinking-water safety"],
        ["Dust-control compliance", "Monthly", "Reduce air pollution"],
        ["Road-condition inspection", "Monthly", "Improve mobility and safety"],
        ["House-damage complaints", "Monthly", "Support compensation decisions"],
        ["Health camp records", "Quarterly", "Monitor common illnesses"],
        ["Local employment data", "Half-yearly", "Track benefit sharing"],
        ["Compensation cases resolved", "Quarterly", "Improve accountability"],
        ["Vulnerability index update", "Annual", "Monitor household risk over time"],
    ], [5, 4, 7])
    doc.add_page_break()


def chapter8(doc):
    heading(doc, "Chapter 8: Conclusion", 1)
    for text in [
        "This project concludes that mining has a deep and multidimensional impact on local communities in Mandamarri. Mining provides employment, supports local economic activities and shapes the identity of the town. However, it also creates environmental, health, livelihood and governance challenges that affect household well-being.",
        "The strongest evidence from the study is the widespread nature of environmental exposure. All surveyed households reported dust pollution, regular noise disturbance and water-quality problems. These findings show that mining-related environmental stress is a settlement-level issue rather than an isolated household complaint.",
        "Livelihood findings show that mining is economically central, with 66.7 percent of households currently connected to mining work. However, employment benefits are not sufficient to remove vulnerability. Many households depend on one earning member, have limited savings, face health risks and lack alternative livelihood options.",
        "Health and financial vulnerability are closely connected. Illness affecting earning capacity, hospitalization and medical expenses can push households into borrowing. Since only 8.3 percent of households reported savings, many families have weak capacity to cope with shocks.",
        "The vulnerability framework shows that risk in Mandamarri is created through the interaction of exposure, sensitivity and adaptive capacity. Exposure is high because dust, noise and water problems are widespread. Sensitivity is increased by debt, dependents, illness, landlessness and income instability. Adaptive capacity is weakened by low savings and absence of mining compensation.",
        "The absence of reported mining compensation is a major concern. Respondents described house cracks, blasting vibration, road damage and water problems, but no household reported receiving mining compensation. This indicates a gap between impact and redress, and it calls for a stronger grievance and compensation mechanism.",
        "The study recommends a people-centred approach to mining development. Safe drinking water, dust control, health camps, road repair, house-damage assessment, local employment priority, social protection, livelihood diversification and community participation should be implemented together.",
        "Mining in Mandamarri should not be understood only as an economic activity. It is also a social and environmental process that shapes daily life. Sustainable mining development will be meaningful only when local communities receive both livelihood opportunities and protection from mining-related harm.",
    ]:
        para(doc, text)
    doc.add_page_break()


def references_appendix(doc, wb):
    heading(doc, "References", 1)
    bullets(doc, [
        "Blaikie, P., Cannon, T., Davis, I., & Wisner, B. (1994). At Risk: Natural Hazards, People's Vulnerability and Disasters. Routledge.",
        "Chambers, R. (1989). Vulnerability, coping and policy. IDS Bulletin, 20(2), 1-7.",
        "Intergovernmental Panel on Climate Change. (2014). Climate Change 2014: Impacts, Adaptation, and Vulnerability. Cambridge University Press.",
        "International Labour Organization. (1999). Social and Labour Issues in Small-scale Mines. ILO.",
        "Turner, B. L. et al. (2003). A framework for vulnerability analysis in sustainability science. Proceedings of the National Academy of Sciences.",
        "Ministry of Coal, Government of India. Annual reports and coal-sector documents.",
        "Singareni Collieries Company Limited. Public information and company documents.",
        "Primary household survey conducted in Mandamarri, Mancherial district, 2026.",
    ])
    doc.add_page_break()
    heading(doc, "Appendix A: Survey Schedule", 1)
    for sheet, title in [
        ("Basic Information", "Basic Information"),
        ("Section_B_Education_Health", "Education and Health"),
        ("Section_C_Livelihood", "Livelihood"),
        ("Section_D_Exposure", "Environmental Exposure"),
        (" Sensitivity (socio-economic v", "Sensitivity"),
        ("Adaptive Capacity", "Adaptive Capacity"),
        ("Shocks, losses & coping strateg", "Shocks, Losses and Coping"),
        ("Perceptions, awareness & partic", "Perceptions and Participation"),
        ("Assets & infrastructure", "Assets and Infrastructure"),
    ]:
        headers, _ = rows_for(wb, sheet)
        heading(doc, title, 2)
        numbered(doc, headers)
    heading(doc, "Appendix B: Sample Household Vulnerability Data", 1)
    _, sheet10 = rows_for(wb, "Sheet10")
    rows = []
    for i, row in enumerate(sheet10[:30], 1):
        rows.append([i, clean(row.get("C1_Main_Livelihood")), f"{float(row.get('EI') or 0):.3f}", f"{float(row.get('SI') or 0):.3f}", f"{float(row.get('ADI') or 0):.3f}", f"{float(row.get('Vulnerability index') or 0):.3f}", clean(row.get("Vulnerability Level"))])
    table(doc, ["S.No.", "Livelihood", "EI", "SI", "ADI", "VI", "Level"], rows, [1.3, 5, 1.7, 1.7, 1.7, 1.7, 2.3])
    heading(doc, "Appendix C: Open-ended Response Extracts", 1)
    _, open_rows = rows_for(wb, "Open-ended")
    extracts = []
    for i, row in enumerate(open_rows[:20], 1):
        extracts.append([i, clean(row.get("J1. What do you think are the main problems your household faces because of mining?"))[:330], clean(row.get("J2. What support or change would help your household the most in reducing these problems?"))[:330]])
    table(doc, ["S.No.", "Problems Reported", "Support Expected"], extracts, [1.2, 7, 7])


def main():
    data = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    doc = Document()
    configure(doc)
    title_pages(doc)
    abstract_toc(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc, data, wb)
    chapter6(doc)
    chapter7(doc)
    chapter8(doc)
    references_appendix(doc, wb)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Finaldoc - Impact of Mining on Local Communities, Mandamarri"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
