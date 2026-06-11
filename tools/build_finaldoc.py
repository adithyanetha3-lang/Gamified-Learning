from pathlib import Path
from collections import Counter
import json
import re

import openpyxl
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path("outputs/pavan_mandamarri")
SOURCE = BASE / "Impact_of_Mining_on_Local_Communities_Mandamarri_Mancherial_Expanded_40_50_Page_Project.docx"
OUT = BASE / "Finaldoc.docx"
ANALYSIS = BASE / "analysis.json"
XLSX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/PavanThesis.0.xlsx")
MAP = Path(r"C:/Users/adith/Downloads/ChatGPT Image Jun 3, 2026, 10_59_57 PM.png")


def clean(v):
    if v is None:
        return ""
    return re.sub(r"\s+", " ", str(v).strip())


def configure(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.28
    p.paragraph_format.space_after = Pt(7)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, header, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def rows_for(wb, sheet):
    ws = wb[sheet]
    rows = list(ws.iter_rows(values_only=True))
    headers = [clean(x) or f"Column {i+1}" for i, x in enumerate(rows[0])]
    out = []
    for row in rows[1:]:
        if any(clean(x) for x in row):
            out.append({headers[i]: row[i] if i < len(row) else None for i in range(len(headers))})
    return headers, out


def freq(rows, key, limit=25):
    counts = Counter(clean(r.get(key)) for r in rows if clean(r.get(key)))
    total = sum(counts.values()) or 1
    return [[k, v, f"{v * 100 / total:.1f}%"] for k, v in counts.most_common(limit)]


def yes_no(rows, key):
    yes = no = other = 0
    for row in rows:
        value = clean(row.get(key)).lower()
        if value.startswith("yes"):
            yes += 1
        elif value.startswith("no"):
            no += 1
        elif value:
            other += 1
    total = yes + no + other or 1
    return [["Yes", yes, f"{yes * 100 / total:.1f}%"], ["No", no, f"{no * 100 / total:.1f}%"], ["Other/Specified", other, f"{other * 100 / total:.1f}%"]]


def add_map(doc):
    add_heading(doc, "Location Map of the Study Area", 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(MAP), width=Inches(6.25))
    cap = doc.add_paragraph("Figure 1: Location map of Mandamarri, Mancherial district, Telangana, India.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    add_para(doc, "The map situates Mandamarri within Mancherial district and Telangana state. Its location near major settlements, road and rail connectivity, and the Godavari River system is important for understanding the local mining economy. The spatial setting also helps explain why mining-related impacts are experienced through transport movement, dust, settlement expansion, water stress and livelihood dependence.")


def add_final_sections(doc, data, wb):
    basic_h, basic = rows_for(wb, "Basic Information")
    edu_h, edu = rows_for(wb, "Section_B_Education_Health")
    liv_h, liv = rows_for(wb, "Section_C_Livelihood")
    exp_h, exp = rows_for(wb, "Section_D_Exposure")
    sens_h, sens = rows_for(wb, " Sensitivity (socio-economic v")
    adap_h, adap = rows_for(wb, "Adaptive Capacity")
    shock_h, shock = rows_for(wb, "Shocks, losses & coping strateg")
    percep_h, percep = rows_for(wb, "Perceptions, awareness & partic")
    asset_h, asset = rows_for(wb, "Assets & infrastructure")
    sheet10_h, sheet10 = rows_for(wb, "Sheet10")

    doc.add_page_break()
    add_heading(doc, "Final Integrated Revisions", 1)
    add_para(doc, "This final revision strengthens the expanded project by integrating the additional Mandamarri methodology documents, the detailed study-area profile, the vulnerability assessment protocol, and the location map. The sections below are written as final submission material and may be read together with the earlier chapters of the project.")

    add_heading(doc, "Elaborated Introduction", 1)
    for heading, paras in [
        ("Mining, Development and Community Life", [
            "Mining is one of the most influential economic activities in northern Telangana. It supports industrial production, generates employment and shapes the growth of settlements such as Mandamarri. In a mining town, the mine is not only a workplace; it becomes a central institution around which transport, housing, small businesses, migration, labour relations and public infrastructure develop.",
            "At the same time, mining produces a complex social reality. The benefits of mining are visible through wages, contract work, transport demand and local markets. The costs are visible through dust, noise, blasting vibration, pressure on water resources, damaged roads, health concerns and uncertainty about compensation. This dual nature makes Mandamarri an important site for studying the impact of mining on local communities.",
            "The present project approaches mining as a multidimensional process. It does not ask only whether mining provides employment. It asks how mining changes everyday life, how households experience environmental pressure, how livelihoods are shaped by mining dependence, and how vulnerable families cope with shocks. This perspective is necessary because mining impact is felt simultaneously in income, health, housing, water, transport and social participation.",
        ]),
        ("Why Mandamarri Matters as a Study Area", [
            "Mandamarri is a prominent coal mining town and mandal in Mancherial district. It forms part of the Godavari Valley coalfield region and has developed historically through coal extraction and related economic activities. The presence of mining has attracted workers from different social backgrounds and has created a settlement pattern closely connected with the coal economy.",
            "The study area is also important because it represents the development paradox common to many mining regions. The same activity that produces employment and infrastructure also creates environmental and social stress. Households living near mining activity may depend on mining income while also facing pollution, water problems, road damage and health risks.",
            "The local economy of Mandamarri is not limited to direct mine employment. It includes contract workers, permanent workers, daily wage labourers, drivers, vendors, small businesses, farmers and service workers. This diversity makes the area suitable for analysing both direct and indirect mining impacts.",
        ]),
        ("Research Gap and Relevance", [
            "Many discussions of mining focus on production, revenue and employment. Fewer studies examine how mining is experienced by households at the settlement level. A household-level study is necessary because community well-being depends on daily access to safe water, clean air, health care, roads, schools, stable income and compensation when damage occurs.",
            "This research fills that gap by using primary survey data from 60 households in Mandamarri. It combines quantitative indicators with open-ended responses, allowing the project to measure vulnerability while also preserving the voices of respondents. The study is therefore relevant for academic understanding as well as local planning.",
            "The project also contributes by applying a Composite Vulnerability Index based on exposure, sensitivity and adaptive capacity. This makes it possible to move beyond general statements and identify the specific drivers of household vulnerability in a mining-affected settlement.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for para in paras:
            add_para(doc, para)
    add_map(doc)

    add_heading(doc, "Revised Study Area Profile: Mandamarri, Mancherial District", 1)
    for heading, paras in [
        ("Location and Administrative Profile", [
            "Mandamarri is located in Mancherial district in the state of Telangana. The district lies in the northern part of Telangana and is widely associated with coal mining, forest resources, river systems and industrial activity. Mandamarri functions as an important mining town within this regional setting.",
            "Administratively, Mandamarri is connected with nearby towns such as Mancherial, Bellampalli, Luxettipet, Chennur and Dandepalli. Road and rail linkages connect the town to the broader coal belt and labour market. These connections are significant because mining-related transport and mobility influence the everyday experience of local residents.",
            "The settlement's position within a mining belt means that its economy, demographic character and environmental conditions are closely connected with the operations of coal extraction. Local households are therefore affected not only by direct mining work but also by the broader changes generated by mining infrastructure.",
        ]),
        ("Physical Geography and Coalfield Setting", [
            "Mandamarri is associated with the Godavari Valley coalfield region, one of the important coal-bearing zones of India. The region is characterized by undulating terrain, dry deciduous vegetation, seasonal climatic conditions and settlement patterns shaped by both natural resources and industrial development.",
            "The coal-bearing geological formations have played a decisive role in Mandamarri's development. Mining activity has altered land use, vegetation, drainage and settlement expansion. Opencast and underground mining practices influence the physical environment in different ways, including dust generation, vibration, water movement and land disturbance.",
            "The proximity of Mandamarri to river systems and transport routes makes environmental management especially important. Water quality, drainage patterns and dust movement are not isolated technical issues; they directly affect household health, agriculture, mobility and quality of life.",
        ]),
        ("Historical Development of Mining", [
            "The history of Mandamarri is closely linked with the growth of coal mining in the Godavari Valley region. As demand for coal increased, mining infrastructure expanded and settlements around the mines developed as labour and service centres.",
            "Mining brought employment opportunities and attracted workers from different communities. Over time, Mandamarri developed a mixed social composition that included mine workers, contract labourers, small traders, farmers, drivers and informal workers. This historical development explains why mining remains central to local identity and livelihood.",
            "In recent decades, technological changes and expansion of opencast mining have changed the nature of mining impacts. While mechanization can increase production, it may also intensify dust, transport movement, land disturbance and concerns about local employment security.",
        ]),
        ("Demographic and Social Character", [
            "Mandamarri's demographic profile reflects its mining history. The settlement includes Scheduled Castes, Scheduled Tribes, Backward Classes, other caste groups and religious minorities. Migration related to mining work has contributed to social diversity and the formation of working-class neighbourhoods.",
            "The household survey shows that the sample includes different caste and livelihood backgrounds. Such diversity is important because vulnerability is not equal across households. Social position, education, land ownership, employment type, health condition and family size influence how people experience mining impacts.",
            "The presence of dependents in many households increases the importance of stable income, health protection and access to public services. When a household has children, elderly members or non-earning dependents, mining-related shocks can quickly become family-level crises.",
        ]),
        ("Economic Structure and Livelihood Pattern", [
            "The economy of Mandamarri is strongly shaped by coal mining and related activities. Singareni Collieries and associated mining operations provide direct and indirect employment. At the same time, the local economy includes agriculture, small shops, transport, private jobs, vending, daily wage labour and service work.",
            "This mixed livelihood structure is important for understanding impact. Mining workers may receive wages but face occupational and contract-related risks. Non-mining households may not receive direct mining benefits but still face environmental and infrastructure impacts. Farmers may experience dust, water and land-related problems, while vendors and drivers depend on mining-linked demand.",
            "The survey indicates that 66.7 percent of households are currently linked to mining work. This demonstrates strong dependence, but dependence does not automatically mean security. Income stability, type of employment, documentation, savings and health condition all influence whether mining dependence becomes a strength or a vulnerability.",
        ]),
        ("Environmental Conditions and Associated Impacts", [
            "The environmental conditions of Mandamarri have been significantly influenced by mining. Dust from excavation, blasting, coal transport and heavy vehicle movement is one of the most visible impacts. Noise and vibration from blasting and machinery also affect the settlement environment.",
            "Water resources are a major concern. Respondents reported water-quality problems and open-ended responses repeatedly mentioned unfiltered water, water scarcity and the need for safe supply. This makes water one of the central issues in the project.",
            "The reduction of vegetation, road damage, air pollution and housing cracks reported by respondents show that mining impact is spread across environmental and infrastructural systems. The study therefore treats environmental impact as a household welfare issue, not only a physical-geography issue.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for para in paras:
            add_para(doc, para)

    add_heading(doc, "Strengthened Methodology from the Mandhamari Methodology Documents", 1)
    for heading, paras in [
        ("Philosophical Orientation", [
            "The methodology follows the view that vulnerability is socially produced and context-specific. Households become vulnerable not only because they are exposed to a hazard, but because their social, economic and institutional conditions shape their ability to respond. This orientation is suitable for Mandamarri because mining impacts are experienced through both physical exposure and household capacity.",
            "The study uses a mixed-method approach. Quantitative survey data provide measurable evidence on household conditions, while qualitative responses and observations explain the meaning of those conditions. This combination is important because vulnerability cannot be understood fully through numbers alone.",
        ]),
        ("IPCC Vulnerability Framework", [
            "The IPCC vulnerability framework is used because it organizes vulnerability into exposure, sensitivity and adaptive capacity. Exposure refers to the degree to which households encounter mining-related stress such as dust, noise, water-quality problems, accidents and proximity to mines. Sensitivity refers to internal household conditions such as illness, dependents, debt, landlessness and unstable income. Adaptive capacity refers to resources such as savings, bank accounts, group membership, health facilities, schools, social protection and compensation.",
            "This framework is appropriate for Mandamarri because the survey shows widespread exposure but different household capacities. Almost every household reported environmental stress, yet households differ in health, debt, income, savings and social support. The framework therefore helps explain why the same mining environment does not affect all households equally.",
        ]),
        ("Pressure and Release Model", [
            "The Pressure and Release model complements the IPCC framework by explaining how vulnerability is produced through root causes, dynamic pressures and unsafe conditions. In Mandamarri, root causes include dependence on mining, limited alternatives, weak compensation and structural inequalities. Dynamic pressures include contract labour, debt, inadequate public services and weak enforcement. Unsafe conditions include dust exposure, water problems, road damage, house cracks and health risks.",
            "Combining the IPCC framework and PAR model allows the project to connect household-level scores with deeper social processes. The index identifies levels of vulnerability, while the PAR model explains why such vulnerability exists.",
        ]),
        ("Research Design and Sampling", [
            "The research design is descriptive and analytical. It is descriptive because it documents the socio-economic, health, livelihood, environmental and infrastructure profile of households. It is analytical because it constructs indices and interprets the drivers of vulnerability.",
            "Purposive sampling was used because the study focuses specifically on mining-affected households in Mandamarri. A total of 60 households were surveyed. The sample includes different livelihoods and social backgrounds, allowing the project to compare patterns of exposure, sensitivity and adaptive capacity.",
        ]),
        ("Data Collection Tools", [
            "The principal tool was a semi-structured household questionnaire. It collected information on basic profile, education, health, livelihood, exposure, sensitivity, adaptive capacity, shocks, perceptions, assets and open-ended experiences.",
            "Personal interviews and field observations supported the survey. These methods were useful for clarifying responses, recording local concerns and understanding the physical environment of the study area. Open-ended responses were especially valuable for identifying issues such as water scarcity, dust, blasting vibration, house cracks, roads and employment expectations.",
        ]),
        ("Indicator Selection and Index Construction", [
            "Indicator selection followed theoretical relevance, data availability, variability and interpretability. Each indicator was assigned to one of three dimensions: Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index.",
            "Indicators were coded numerically and normalized to make them comparable. Positive indicators increase vulnerability when their value rises, while negative indicators are reversed so that higher normalized values consistently indicate higher vulnerability. The three sub-indices were then aggregated into the Composite Vulnerability Index.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for para in paras:
            add_para(doc, para)

    add_table(doc, ["Step", "Methodological action in this study", "Purpose"], [
        ["1", "Defined research aim, study area and household-level mining vulnerability", "Clarifies the scope of assessment"],
        ["2", "Selected household survey and mixed-method assessment", "Captures both measurable and experiential impacts"],
        ["3", "Adopted Tier-2 indicator-based composite index", "Creates a structured vulnerability measurement"],
        ["4", "Defined sector as mining-affected local community in Mandamarri", "Connects household vulnerability with mining context"],
        ["5", "Selected indicators across exposure, sensitivity and adaptive capacity", "Ensures multidimensional coverage"],
        ["6", "Coded household responses numerically", "Prepares survey data for analysis"],
        ["7", "Normalized indicators using min-max logic", "Makes indicators comparable across scales"],
        ["8", "Assigned weights and grouped indicators", "Reflects relative contribution to vulnerability"],
        ["9", "Computed EI, SI, ADI and overall CVI", "Measures household vulnerability"],
        ["10", "Presented results through tables, charts and maps", "Improves interpretation and communication"],
        ["11", "Classified households into vulnerability levels", "Identifies low, medium and high vulnerability groups"],
        ["12", "Interpreted drivers using IPCC and PAR models", "Connects numbers with social explanation"],
    ], [1.5, 8, 6])

    add_table(doc, ["Indicator type", "Treatment", "Example from study"], [
        ["Positive indicator", "Higher value means higher vulnerability", "Dust pollution, debt, illness, accident experience"],
        ["Negative indicator", "Higher original value means lower vulnerability, so score is reversed", "Education, stable income, savings, social protection"],
        ["Binary indicator", "Coded as presence or absence", "Health insurance, bank account, mining compensation"],
        ["Ordinal indicator", "Ordered categories converted into scores", "Education level, income stability, housing type"],
    ], [4, 6, 6])

    add_heading(doc, "Elaborated Results", 1)
    for heading, paras in [
        ("Demographic Results", [
            "The demographic results show that the surveyed households represent Mandamarri's working-class and mining-influenced social structure. The sample contains different caste groups, religions, housing types and household sizes. The average respondent age is 41.65 years, showing that most respondents are within working and family-responsibility age groups.",
            "Male respondents form 90 percent of the sample and female respondents form 10 percent. This reflects the field survey pattern and the gendered nature of public household representation in many mining settlements. However, the impacts described in the study affect the entire household, including women, children and elderly members.",
            "Household size averages 3.4 persons, and 85 percent of households reported dependents. This is important because dependents increase sensitivity to livelihood shocks. If one earning member faces illness, accident or job interruption, the impact spreads to all dependent members.",
        ]),
        ("Education and Health Results", [
            "Education levels show that 43.3 percent of respondents have higher secondary education and 23.3 percent are graduates. This indicates a meaningful human-capital base in Mandamarri. Technical and vocational training is also present among many households, which can support employment diversification if proper opportunities are created.",
            "Health results reveal important vulnerability. Illness affected earning capacity in 38.3 percent of households, hospitalization was reported by 35 percent, and only 35 percent reported health insurance. These figures show that health shocks can directly affect income security.",
            "Mining-affected communities need health systems that address both general and occupational risks. Respiratory problems, skin disease, back pain, accident risk, stress and chronic illness can interact with livelihood insecurity. Periodic health camps and insurance enrolment should therefore be treated as essential interventions.",
        ]),
        ("Livelihood Results", [
            "Livelihood data show that mining is central to the local economy. Around 66.7 percent of households are currently linked to mining work. Mining workers, contract workers, coal cutters, trolley drivers, security guards, daily wage labourers, farmers, vendors and private job workers form the occupational profile.",
            "The average number of economically active members is only 1.1 per household. This indicates dependence on one main earner. Such dependence increases household vulnerability because income interruption, illness or job loss cannot easily be absorbed by other earning members.",
            "Although 60 percent of households reported stable income, 40 percent reported seasonal, unstable or sometimes low income. This shows that mining-town livelihoods are not uniformly secure. Contract work, informal work and small business are particularly sensitive to market and employment fluctuations.",
        ]),
        ("Environmental Exposure Results", [
            "Environmental exposure is the strongest finding of the project. All 60 households reported dust or airborne pollution, regular noise disturbance and water-quality problems. This indicates that environmental stress is widespread and forms the everyday background of life in Mandamarri.",
            "The open-ended responses strongly support the quantitative results. Respondents mentioned unfiltered water, summer water scarcity, dust from opencast activity, damaged roads, blasting vibration and cracks in houses. These responses show that environmental impacts are experienced through concrete household problems.",
            "Accidents were reported by 13.3 percent of households, while agriculture or livestock impact was reported by a smaller section. Even when a specific impact affects fewer households, it can be severe for those families. Therefore, policy should not ignore low-frequency but high-loss impacts.",
        ]),
        ("Sensitivity Results", [
            "Sensitivity indicators show why some households are more affected than others. Sixty percent of households reported no land, 35 percent reported debt and 85 percent reported dependents. Landlessness reduces asset security, debt increases financial pressure and dependents increase household responsibility.",
            "Livestock ownership is low, which limits livelihood diversification. Although toilet access and tap-water access are relatively high, respondents still reported water-quality concerns. This distinction is important: access to a facility does not automatically mean the service is safe, reliable or sufficient.",
            "Sensitivity is therefore not a single problem. It is a combination of landlessness, family responsibility, debt, health insecurity, livelihood dependence and weak alternatives.",
        ]),
        ("Adaptive Capacity Results", [
            "Adaptive capacity presents a mixed picture. All households reported bank accounts and access to a primary school and PHC. Group membership is also present among many households. These are positive institutional and social-capital indicators.",
            "However, only 8.3 percent reported savings and no household reported mining compensation. This is a major weakness. Savings are the first household-level buffer during crisis, and compensation is essential when mining causes damage.",
            "The adaptive-capacity results show that public services exist but are not enough. Households need stronger financial resilience, health protection, compensation mechanisms and alternative livelihood opportunities.",
        ]),
        ("Perception and Open-ended Results", [
            "Perception data show that residents understand mining as a serious local issue. Respondents repeatedly demanded safe water, filter beds, regular transport, road improvement, local employment, compensation for house damage and access to better amenities.",
            "These demands are practical and development-oriented. They show that the community is not simply opposing mining; rather, residents want mining to become more accountable and beneficial to the local population.",
            "The open-ended responses are important because they convert numerical findings into lived experience. They show the human side of dust, water, noise, road damage and employment insecurity.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for para in paras:
            add_para(doc, para)

    add_table(doc, ["Theme", "Key result", "Interpretation"], [
        ["Mining livelihood", "66.7% currently linked to mining", "Mining is economically central but creates dependence"],
        ["Dust and noise", "100% reported dust and regular noise", "Exposure is settlement-wide"],
        ["Water quality", "100% reported water-quality problem", "Safe water is a priority intervention"],
        ["Health", "38.3% reported illness affecting earning", "Health shocks reduce livelihood security"],
        ["Dependents", "85% have dependents", "Family responsibility increases sensitivity"],
        ["Savings", "Only 8.3% have savings", "Financial coping capacity is weak"],
        ["Compensation", "0% reported mining compensation", "Governance and accountability gap is visible"],
    ], [4, 5, 7])

    add_heading(doc, "Elaborated Discussion", 1)
    for heading, paras in [
        ("Mining as Both Opportunity and Vulnerability", [
            "The discussion of Mandamarri must begin with the dual character of mining. Mining provides employment, wages, transport demand and local business opportunities. At the same time, it produces environmental pressure and social insecurity. This means that mining cannot be judged only as beneficial or harmful; it must be evaluated through the distribution of benefits and costs.",
            "Households directly connected to mining may receive income but still experience pollution and health risks. Households not directly employed in mining may experience the environmental costs without receiving secure benefits. This uneven distribution is one of the central issues in mining-affected development.",
            "The project findings show that mining dependence can become a form of vulnerability when alternative livelihoods are weak. If the community depends heavily on one industry, any instability in employment or health can create household-level crisis.",
        ]),
        ("Environmental Justice Perspective", [
            "From an environmental justice perspective, the key question is who bears the burden of mining and who receives the benefits. In Mandamarri, households report widespread dust, noise and water-quality problems. These are collective burdens experienced by the settlement.",
            "The absence of mining compensation despite reports of house cracks and damage indicates a gap between impact and redress. Environmental justice requires not only pollution control but also recognition, participation and compensation.",
            "A fair approach would include transparent environmental monitoring, public disclosure of results, community grievance forums and compensation for verified damage.",
        ]),
        ("Vulnerability Drivers", [
            "The Composite Vulnerability Index helps identify drivers of household vulnerability. Exposure is high because environmental pressures are common. Sensitivity is shaped by dependents, debt, illness, landlessness and unstable income. Adaptive capacity is weakened by lack of savings, limited alternatives and absence of compensation.",
            "The PAR model helps explain this pattern. Root causes include dependence on coal mining and weak local alternatives. Dynamic pressures include contract work, debt, limited health protection and weak grievance mechanisms. Unsafe conditions include dust, noise, water problems, damaged roads and housing cracks.",
            "The discussion therefore shows that household vulnerability is produced by the interaction of environmental and social factors. Reducing vulnerability requires action on both sides.",
        ]),
        ("Policy Implications", [
            "The findings have clear policy implications. Pollution control alone will not solve the problem if households remain economically insecure. Employment schemes alone will not solve the problem if water, dust, roads and health risks are ignored.",
            "A comprehensive mining-area development plan should combine environmental management, social protection, local employment, health care, compensation, skill training and community participation.",
            "The most urgent interventions are safe drinking water, dust reduction, road repair, health camps, house-damage verification and local job priority. These measures respond directly to the strongest findings of the survey.",
        ]),
    ]:
        add_heading(doc, heading, 2)
        for para in paras:
            add_para(doc, para)

    add_heading(doc, "Elaborated Conclusion", 1)
    for para in [
        "This final project concludes that the impact of mining on local communities in Mandamarri is multidimensional and deeply embedded in everyday life. Mining provides employment and supports the local economy, but it also creates environmental stress, health concerns, infrastructure damage and household vulnerability.",
        "The most powerful evidence comes from environmental exposure. All surveyed households reported dust pollution, regular noise disturbance and water-quality problems. These findings indicate that mining-related environmental stress is not limited to a few households but is experienced widely across the study area.",
        "Livelihood results show that mining remains economically central, with 66.7 percent of households linked to mining work. However, the low average number of earning members, unstable income for many households, debt, illness and lack of savings show that employment alone does not guarantee security. Mining dependence must therefore be accompanied by livelihood protection and diversification.",
        "The vulnerability framework demonstrates that risk is created through the interaction of exposure, sensitivity and adaptive capacity. Mandamarri households face high environmental exposure, while sensitivity is increased by dependents, debt, landlessness, illness and unstable income. Adaptive capacity is weakened by low savings and absence of compensation.",
        "The absence of reported mining compensation is a major governance concern. When households experience dust, water problems, road damage, blasting vibration and house cracks, a fair and transparent redress system is essential. Without such a system, mining-led development can produce resentment and social injustice.",
        "The study recommends safe water supply, filtration, dust control, road repair, local employment priority, health camps, insurance coverage, house-damage assessment, compensation mechanisms and community participation. These interventions should be implemented together because the problems are interconnected.",
        "In conclusion, Mandamarri requires a people-centred mining development approach. Mining should not only extract resources and create employment; it should also protect the health, housing, water access, environment and dignity of the people who live closest to the mines. Sustainable development in Mandamarri will be possible only when local communities receive both economic opportunities and protection from mining-related harm.",
    ]:
        add_para(doc, para)

    add_heading(doc, "Additional Final Tables for Submission", 1)
    add_heading(doc, "Table: Education Distribution", 2)
    add_table(doc, ["Education", "Households", "Percentage"], freq(edu, "B1_Education"), [6, 4, 4])
    add_heading(doc, "Table: Livelihood Distribution", 2)
    add_table(doc, ["Livelihood", "Households", "Percentage"], freq(liv, "C1_Main_Livelihood", 30), [7, 4, 4])
    add_heading(doc, "Table: Environmental Exposure Indicators", 2)
    add_table(doc, ["Indicator", "Households", "Percentage"], [
        ["Dust or airborne pollution", "60", "100.0%"],
        ["Regular noise disturbance", "60", "100.0%"],
        ["Water-quality problem", "60", "100.0%"],
        ["Accident experience", "8", "13.3%"],
        ["Agriculture/livestock affected", "5", "8.3%"],
    ], [7, 4, 4])
    add_heading(doc, "Table: Sensitivity and Adaptive Capacity Indicators", 2)
    add_table(doc, ["Indicator", "Finding", "Meaning"], [
        ["No land", "36 households (60.0%)", "Weak asset base"],
        ["Debt", "21 households (35.0%)", "Financial pressure"],
        ["Dependents", "51 households (85.0%)", "Higher household responsibility"],
        ["Savings", "5 households (8.3%)", "Weak shock-coping capacity"],
        ["Bank account", "60 households (100.0%)", "Institutional access exists"],
        ["Mining compensation", "0 households (0.0%)", "Major redress gap"],
    ], [5, 5, 6])
    add_heading(doc, "Table: Vulnerability-Level Classification", 2)
    add_table(doc, ["Vulnerability Level", "Households", "Percentage"], freq(sheet10, "Vulnerability Level"), [6, 4, 4])


def update_title(doc):
    for p in doc.paragraphs[:12]:
        if "IMPACT OF MINING ON LOCAL COMMUNITIES" in p.text:
            p.text = ""
            r = p.add_run("IMPACT OF MINING ON LOCAL COMMUNITIES\nMANDAMARRI, MANCHERIAL DISTRICT, TELANGANA")
            r.bold = True
            r.font.size = Pt(18)
            r.font.color.rgb = RGBColor(31, 78, 121)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in doc.paragraphs[:18]:
        if "Final Research Project Draft" in p.text:
            p.text = ""
            r = p.add_run("Finaldoc - Ready to Submit Project Report")
            r.bold = True
            r.font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    data = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    doc = Document(SOURCE)
    configure(doc)
    update_title(doc)
    add_final_sections(doc, data, wb)
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Finaldoc - Impact of Mining on Local Communities, Mandamarri"
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
