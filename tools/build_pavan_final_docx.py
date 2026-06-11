from pathlib import Path
from collections import Counter
import json
import re

import openpyxl
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path("outputs/pavan_mandamarri")
ANALYSIS = BASE / "analysis.json"
XLSX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/PavanThesis.0.xlsx")
OUT = BASE / "Impact_of_Mining_on_Local_Communities_Mandamarri_Mancherial_Final_Draft.docx"


def pct(n, d=60):
    return round((n * 100.0) / d, 1) if d else 0


def clean(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip())


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color="FFFFFF")
        set_cell_shading(hdr[i], "1F4E79")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def freq_sheet(wb, sheet_name, col_name):
    ws = wb[sheet_name]
    rows = list(ws.iter_rows(values_only=True))
    headers = [clean(v) for v in rows[0]]
    idx = headers.index(col_name)
    c = Counter(clean(row[idx]) for row in rows[1:] if clean(row[idx]))
    return c


def read_sheet10(wb):
    ws = wb["Sheet10"]
    rows = list(ws.iter_rows(values_only=True))
    headers = [clean(v) or f"Column {i+1}" for i, v in enumerate(rows[0])]
    data = []
    for row in rows[1:]:
        if any(clean(v) for v in row):
            data.append({headers[i]: row[i] if i < len(row) else None for i in range(len(headers))})
    return data


def chart_bar(path, title, labels, values, color="#2F6F73"):
    width, height = 1300, 760
    margin_l, margin_r, margin_t, margin_b = 120, 60, 100, 190
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34, bold=True)
    axis_font = font(22)
    label_font = font(18)
    value_font = font(20, bold=True)
    draw.text((width // 2, 35), title, fill="#1F4E79", font=title_font, anchor="ma")
    plot_w = width - margin_l - margin_r
    plot_h = height - margin_t - margin_b
    max_val = max(values) if values else 1
    for i in range(6):
        y = margin_t + plot_h - int(plot_h * i / 5)
        draw.line((margin_l, y, width - margin_r, y), fill="#D9E2F3", width=1)
        tick = round(max_val * i / 5)
        draw.text((margin_l - 14, y), str(tick), fill="#555555", font=label_font, anchor="rm")
    draw.line((margin_l, margin_t, margin_l, margin_t + plot_h), fill="#555555", width=2)
    draw.line((margin_l, margin_t + plot_h, width - margin_r, margin_t + plot_h), fill="#555555", width=2)
    n = len(values)
    gap = 28
    bar_w = max(42, int((plot_w - gap * (n + 1)) / max(n, 1)))
    for idx, (label, value) in enumerate(zip(labels, values)):
        x0 = margin_l + gap + idx * (bar_w + gap)
        x1 = x0 + bar_w
        bar_h = int(plot_h * value / max_val) if max_val else 0
        y0 = margin_t + plot_h - bar_h
        y1 = margin_t + plot_h
        draw.rectangle((x0, y0, x1, y1), fill=color)
        draw.text(((x0 + x1) // 2, y0 - 12), str(value), fill="#222222", font=value_font, anchor="mb")
        wrapped = wrap_label(label, 16)
        draw.multiline_text(((x0 + x1) // 2, y1 + 18), wrapped, fill="#222222", font=label_font, anchor="ma", align="center", spacing=3)
    draw.text((34, margin_t + plot_h // 2), "Households", fill="#555555", font=axis_font, anchor="mm")
    img.save(path)


def chart_pie(path, title, labels, values):
    colors = ["#1F4E79", "#70AD47", "#FFC000", "#C00000", "#7030A0"]
    width, height = 1200, 760
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(34, bold=True)
    label_font = font(23)
    small_font = font(20)
    draw.text((width // 2, 35), title, fill="#1F4E79", font=title_font, anchor="ma")
    box = (110, 130, 640, 660)
    total = sum(values) or 1
    start = -90
    for i, value in enumerate(values):
        extent = value * 360 / total
        draw.pieslice(box, start, start + extent, fill=colors[i % len(colors)], outline="white", width=3)
        mid = start + extent / 2
        start += extent
        percent = value * 100 / total
        import math
        cx, cy, r = 370, 395, 165
        x = cx + r * math.cos(math.radians(mid))
        y = cy + r * math.sin(math.radians(mid))
        draw.text((x, y), f"{percent:.1f}%", fill="white", font=small_font, anchor="mm")
    y = 160
    for i, (label, value) in enumerate(zip(labels, values)):
        draw.rectangle((720, y, 760, y + 40), fill=colors[i % len(colors)])
        draw.text((780, y + 20), f"{label}: {value} households", fill="#222222", font=label_font, anchor="lm")
        y += 64
    img.save(path)


def font(size, bold=False):
    candidates = [
        r"C:/Windows/Fonts/calibrib.ttf" if bold else r"C:/Windows/Fonts/calibri.ttf",
        r"C:/Windows/Fonts/arialbd.ttf" if bold else r"C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            pass
    return ImageFont.load_default()


def wrap_label(label, width):
    words = str(label).split()
    lines = []
    current = ""
    for word in words:
        trial = (current + " " + word).strip()
        if len(trial) <= width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return "\n".join(lines[:3])


def add_picture(doc, image_path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def open_ended_themes(wb):
    ws = wb["Open-ended"]
    rows = list(ws.iter_rows(values_only=True))[1:]
    problems = " ".join(clean(r[0]).lower() for r in rows)
    support = " ".join(clean(r[1]).lower() for r in rows)
    problem_map = {
        "Water scarcity and unsafe water": ["water scarcity", "unfiltered water", "water"],
        "Dust and air pollution": ["dust", "air pollution"],
        "Blasting vibration and house cracks": ["cracks", "vibration", "blasting"],
        "Damaged roads and transport gaps": ["road", "transport"],
        "Limited jobs and local employment": ["job", "employment"],
        "Health problems": ["skin", "allergy", "disease", "health"],
    }
    support_map = {
        "Regular safe water supply/filter bed": ["water", "filterbed", "godavari"],
        "Local job reservation/employment priority": ["job", "employment", "reservation"],
        "Compensation for house damage": ["compensation", "damaged houses", "infrastructure loss"],
        "Transport, roads, school and park": ["transport", "road", "school", "park"],
        "Basic amenities for land/settlement": ["amenities", "amenties", "land"],
    }
    out = []
    for label, keys in problem_map.items():
        out.append((label, sum(problems.count(k) for k in keys)))
    supp = []
    for label, keys in support_map.items():
        supp.append((label, sum(support.count(k) for k in keys)))
    return sorted(out, key=lambda x: x[1], reverse=True), sorted(supp, key=lambda x: x[1], reverse=True)


def configure_document(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.3)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def main():
    BASE.mkdir(parents=True, exist_ok=True)
    data = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    sheet10 = read_sheet10(wb)
    vuln_levels = Counter(clean(r.get("Vulnerability Level")) for r in sheet10 if clean(r.get("Vulnerability Level")))
    problems, supports = open_ended_themes(wb)

    chart1 = BASE / "vulnerability_levels.png"
    chart2 = BASE / "environmental_exposure.png"
    chart3 = BASE / "livelihood_sources.png"
    chart_pie(chart1, "Composite Vulnerability Levels", list(vuln_levels.keys()), list(vuln_levels.values()))
    chart_bar(
        chart2,
        "Environmental Exposure Reported by Households",
        ["Dust", "Noise", "Water", "Accident", "Agri/Livestock"],
        [
            data["environmental_exposure"]["dust"]["yes"],
            60,
            60,
            data["environmental_exposure"]["accident"]["yes"],
            5,
        ],
        "#C0504D",
    )
    top_income = data["livelihood"]["income_source"][:6]
    chart_bar(chart3, "Main Income Sources", [x["value"] for x in top_income], [x["count"] for x in top_income], "#4F81BD")

    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("IMPACT OF MINING ON LOCAL COMMUNITIES\nMANDAMARRI, MANCHERIAL DISTRICT, TELANGANA")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(31, 78, 121)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Final Research Project Draft").bold = True
    doc.add_paragraph()
    details = [
        ["Submitted by", "Durgam Pavan"],
        ["Study Area", "Mandamarri, Mancherial District, Telangana"],
        ["Sample Size", "60 households"],
        ["Method", "Household survey, field observations, open-ended responses and Composite Vulnerability Index"],
        ["Date", "June 2026"],
    ]
    add_table(doc, ["Particular", "Details"], details, [4, 11])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Ready-to-present draft prepared from the field survey workbook.").italic = True
    doc.add_page_break()

    add_heading(doc, "Declaration", 1)
    add_para(doc, "I hereby declare that this project report titled Impact of Mining on Local Communities: A Study of Mandamarri, Mancherial District is based on primary household survey data, field observations and secondary conceptual understanding. The interpretations, tables and conclusions presented in this report are prepared for academic purposes.")
    doc.add_page_break()

    add_heading(doc, "Acknowledgement", 1)
    add_para(doc, "I express my sincere gratitude to the respondents of Mandamarri who shared their time, experiences and concerns about mining, livelihood, health, water, transport and local development. Their participation made this study possible. I also thank my teachers, classmates and family members for their encouragement and guidance during the preparation of this project.")
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_para(doc, "Mining has contributed to employment and industrial growth in the coal belt of Mancherial district, but it has also created serious pressures on local communities. This study examines the impact of mining on 60 households in Mandamarri. The research focuses on livelihood dependence, environmental exposure, health problems, social sensitivity, adaptive capacity and household-level vulnerability.")
    add_para(doc, "The findings show that environmental exposure is widespread: all surveyed households reported dust pollution, regular noise disturbance and water-quality problems. Mining-related livelihoods remain important, with 66.7 percent of households currently connected to mining work. At the same time, the community faces weak savings, high dependence on one earning member, debt among 35 percent of households, and a lack of mining compensation. The Composite Vulnerability Index indicates that vulnerability is not produced by one factor alone; it is created by the interaction of environmental exposure, livelihood insecurity, health risks and limited adaptive capacity.")
    add_para(doc, "The study recommends local employment priority, safe drinking water interventions, compensation for damaged houses, health screening, pollution control, transport improvement and stronger community participation in mining-related decision-making.")
    add_para(doc, "Keywords: Mining impact, local communities, Mandamarri, Mancherial, coal mining, vulnerability, livelihood, environmental exposure.")
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    contents = [
        "Chapter 1: Introduction",
        "Chapter 2: Review of Literature and Conceptual Framework",
        "Chapter 3: Research Methodology",
        "Chapter 4: Study Area and Socio-Economic Profile",
        "Chapter 5: Results and Data Analysis",
        "Chapter 6: Composite Vulnerability Index",
        "Chapter 7: Major Findings and Recommendations",
        "Chapter 8: Conclusion",
        "References",
        "Appendices",
    ]
    add_numbered(doc, contents)
    doc.add_page_break()

    add_heading(doc, "Chapter 1: Introduction", 1)
    add_heading(doc, "1.1 Background of the Study", 2)
    add_para(doc, "Mining is an important economic activity in Telangana's coal belt. It provides direct and indirect employment, supports transport and small business activities, and connects local settlements to wider industrial growth. However, mining also changes the everyday life of people living near mining areas. Dust, noise, blasting vibration, heavy vehicle movement, pressure on water resources and uncertainty in employment can create long-term social and environmental stress.")
    add_para(doc, "Mandamarri in Mancherial district is a mining-influenced area where households depend on different forms of work, including mining, daily wage labour, agriculture, small business, driving, security work and private jobs. This study examines how mining affects these local communities across livelihood, health, environment, assets, infrastructure, coping capacity and perceptions.")
    add_heading(doc, "1.2 Problem Statement", 2)
    add_para(doc, "The central problem is that mining-led development creates both opportunities and risks. While mining creates employment for some households, the benefits are unevenly distributed. Many households experience environmental pollution, water problems, damaged roads, house cracks from blasting vibrations, weak compensation systems and limited alternative livelihoods. Therefore, the study asks whether mining improves local well-being or increases vulnerability among affected households.")
    add_heading(doc, "1.3 Aim of the Study", 2)
    add_para(doc, "The aim of the study is to assess the impact of mining on local communities in Mandamarri, Mancherial district, with special attention to socio-economic conditions, environmental exposure, health risks, adaptive capacity and household vulnerability.")
    add_heading(doc, "1.4 Objectives", 2)
    add_bullets(doc, [
        "To study the socio-economic profile of households living in the mining-affected area.",
        "To identify major environmental problems experienced by local communities due to mining.",
        "To examine the role of mining in household livelihood and income stability.",
        "To assess household health issues, debt, assets and coping capacity.",
        "To construct and interpret a Composite Vulnerability Index using exposure, sensitivity and adaptive capacity indicators.",
        "To suggest practical measures for reducing mining-related risks in Mandamarri.",
    ])
    add_heading(doc, "1.5 Research Questions", 2)
    add_bullets(doc, [
        "What are the major impacts of mining on local households in Mandamarri?",
        "How does mining affect livelihood security, health and environmental conditions?",
        "Which households are more vulnerable and why?",
        "What interventions are required to reduce mining-related vulnerability?",
    ])
    add_heading(doc, "1.6 Scope of the Study", 2)
    add_para(doc, "The study is limited to 60 surveyed households in Mandamarri. It focuses on local-level impacts and does not claim to represent every mining area in Telangana. The report is useful for understanding household-level vulnerability and community needs in a specific mining-affected settlement.")

    add_heading(doc, "Chapter 2: Review of Literature and Conceptual Framework", 1)
    add_para(doc, "Studies on mining and local communities show that extractive industries often produce a development paradox. Mining can generate jobs and public revenue, but the same activity may also create environmental degradation, land-use change, occupational health risks, displacement pressures and unequal access to benefits. Communities close to mines often bear the immediate burden of dust, noise, water stress and infrastructure damage.")
    add_para(doc, "The concept of vulnerability helps explain this situation. Vulnerability is not only exposure to a hazard; it also includes the sensitivity of people and the capacity available to respond. A household with low savings, illness, debt, unstable work and weak access to compensation may suffer more from the same mining exposure than a household with secure income and stronger assets.")
    add_heading(doc, "2.1 Conceptual Framework", 2)
    add_para(doc, "This project uses a vulnerability framework built around three linked dimensions: exposure, sensitivity and adaptive capacity. Exposure refers to the direct environmental and physical contact with mining impacts. Sensitivity refers to social and economic conditions that make households more likely to be harmed. Adaptive capacity refers to the resources, services and institutions that help households cope and recover.")
    add_table(doc, ["Dimension", "Meaning in this study", "Selected indicators"], [
        ["Exposure", "Immediate mining-related pressure on households and environment", "Distance from mine, dust, noise, accident, water quality, agriculture/livestock impact"],
        ["Sensitivity", "Household conditions that increase risk", "Education, illness, dependents, debt, land ownership, income stability, social marginalization"],
        ["Adaptive Capacity", "Ability to cope with and recover from stress", "Savings, bank account, group membership, social protection, compensation, alternative livelihood, health and school access"],
    ], [4, 5.5, 6])

    add_heading(doc, "Chapter 3: Research Methodology", 1)
    add_heading(doc, "3.1 Research Design", 2)
    add_para(doc, "The study follows a descriptive and analytical research design. Quantitative information was collected through a structured household survey, while qualitative insights were gathered through open-ended responses about mining-related problems and expected support.")
    add_heading(doc, "3.2 Data Source and Sample", 2)
    add_para(doc, "The primary dataset consists of 60 households from Mandamarri. The Excel workbook includes sections on basic information, education and health, livelihood, environmental exposure, socio-economic sensitivity, adaptive capacity, shocks and coping strategies, perceptions, assets and open-ended responses.")
    add_table(doc, ["Item", "Description"], [
        ["Study area", "Mandamarri, Mancherial district, Telangana"],
        ["Sample size", "60 households"],
        ["Data collection tool", "Structured household schedule"],
        ["Unit of analysis", "Household"],
        ["Analytical method", "Frequency analysis, descriptive statistics and Composite Vulnerability Index"],
    ], [5, 10])
    add_heading(doc, "3.3 Composite Vulnerability Index Method", 2)
    add_para(doc, "The Composite Vulnerability Index was developed by combining Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index. The score allows comparison across households and helps classify vulnerability levels as low, medium or high. Higher values indicate higher vulnerability.")
    add_table(doc, ["Index", "Mean", "Minimum", "Maximum", "Interpretation"], [
        ["Exposure Index", data["potential_index_sheets"][-1]["numeric_stats"]["EI"]["mean"], data["potential_index_sheets"][-1]["numeric_stats"]["EI"]["min"], data["potential_index_sheets"][-1]["numeric_stats"]["EI"]["max"], "Environmental pressure is consistently present"],
        ["Sensitivity Index", data["potential_index_sheets"][-1]["numeric_stats"]["SI"]["mean"], data["potential_index_sheets"][-1]["numeric_stats"]["SI"]["min"], data["potential_index_sheets"][-1]["numeric_stats"]["SI"]["max"], "Household vulnerability varies by income, health and dependents"],
        ["Adaptive Capacity Deficit", data["potential_index_sheets"][-1]["numeric_stats"]["ADI"]["mean"], data["potential_index_sheets"][-1]["numeric_stats"]["ADI"]["min"], data["potential_index_sheets"][-1]["numeric_stats"]["ADI"]["max"], "Weak savings and compensation reduce coping ability"],
        ["Overall Vulnerability Index", data["potential_index_sheets"][-1]["numeric_stats"]["Vulnerability index"]["mean"], data["potential_index_sheets"][-1]["numeric_stats"]["Vulnerability index"]["min"], data["potential_index_sheets"][-1]["numeric_stats"]["Vulnerability index"]["max"], "Moderate average vulnerability with household-level variation"],
    ], [4, 2.2, 2.2, 2.2, 5.2])

    add_heading(doc, "Chapter 4: Study Area and Socio-Economic Profile", 1)
    add_para(doc, "Mandamarri is located in Mancherial district and is influenced by coal mining activities. The settlement reflects a mixed livelihood structure where mining work, daily wage labour, agriculture, driving, small business and private jobs coexist. The field responses show that mining is not only an economic activity but also a major force shaping water access, roads, housing, health and local employment expectations.")
    add_heading(doc, "4.1 Demographic Profile", 2)
    add_table(doc, ["Variable", "Major finding"], [
        ["Gender", "Male respondents: 54 (90.0%); female respondents: 6 (10.0%)"],
        ["Age", f"Mean age: {data['demographics']['age']['mean']} years; median age: {data['demographics']['age']['median']} years"],
        ["Household size", f"Mean household size: {data['demographics']['household_size']['mean']} persons"],
        ["Religion", "Hindu households form the majority of the sample"],
        ["Housing", "Most households live in semi-pucca houses"],
    ], [5, 10])
    add_heading(doc, "4.2 Education and Health", 2)
    add_table(doc, ["Indicator", "Result"], [
        ["Higher secondary education", "26 households/respondents (43.3%)"],
        ["Graduation", "14 households/respondents (23.3%)"],
        ["Vocational/technical training", "40 households (66.7%)"],
        ["Illness affecting earning", "23 households (38.3%)"],
        ["Hospitalization", "21 households (35.0%)"],
        ["Health insurance", "21 households (35.0%)"],
    ], [6, 9])
    add_para(doc, "The education profile shows that a sizeable section has reached higher secondary level, but health insecurity remains an important issue. Chronic illnesses and hospitalization expenses reduce the ability of households to cope with mining-related pressures.")

    add_heading(doc, "Chapter 5: Results and Data Analysis", 1)
    add_heading(doc, "5.1 Livelihood and Income", 2)
    add_para(doc, "The livelihood pattern shows strong direct and indirect dependence on mining. Around 66.7 percent of households reported current connection with mining work. Mining workers, contract workers, coal cutters, trolley drivers, security guards, daily wage labourers, farmers and small vendors together form the livelihood structure of Mandamarri.")
    add_picture(doc, chart3, "Figure 1: Main income sources among surveyed households")
    add_table(doc, ["Livelihood indicator", "Finding"], [
        ["Currently connected to mining", "40 households (66.7%)"],
        ["Past mining connection", "3 households (5.0%)"],
        ["Average economically active members", "1.1 per household"],
        ["Stable income", "36 households (60.0%)"],
        ["Seasonal/not stable income", "24 households (40.0%)"],
        ["Formal documentation", "46 households (76.7%)"],
    ], [6, 9])
    add_para(doc, "Although a majority reported stable income, the low average number of earning members indicates dependence on one main earner. This makes households vulnerable when illness, job loss, accident or income interruption occurs.")

    add_heading(doc, "5.2 Environmental Exposure", 2)
    add_para(doc, "Environmental exposure is the strongest and most uniform finding of the study. Every surveyed household reported dust or airborne pollution, regular noise disturbance and water-quality problems. Open-ended responses repeatedly mention unfiltered water, water scarcity in summer, damaged roads, blasting vibration, cracks in houses and dust from open-cast activity.")
    add_picture(doc, chart2, "Figure 2: Environmental exposure reported by households")
    add_table(doc, ["Exposure indicator", "Households reporting impact", "Percentage"], [
        ["Dust/airborne pollution", "60", "100.0%"],
        ["Regular noise disturbance", "60", "100.0%"],
        ["Water-quality problem", "60", "100.0%"],
        ["Accident experience", "8", "13.3%"],
        ["Agriculture/livestock affected", "5", "8.3%"],
        ["Displacement/relocation", "0", "0.0%"],
    ], [5.5, 5, 3])
    add_heading(doc, "5.3 Sensitivity Conditions", 2)
    add_para(doc, "Sensitivity refers to household conditions that increase the effect of mining-related stress. In Mandamarri, landlessness, dependents, debt, illness and lack of savings are important sensitivity factors.")
    add_table(doc, ["Sensitivity indicator", "Finding"], [
        ["No land", "36 households (60.0%)"],
        ["Dependents present", "51 households (85.0%)"],
        ["Debt", "21 households (35.0%)"],
        ["Livestock ownership", "5 households (8.3%)"],
        ["Private toilet facility", "60 households (100.0%)"],
        ["Tap water in house", "53 households (88.3%), but quality concerns remain"],
    ], [6, 9])
    add_heading(doc, "5.4 Adaptive Capacity", 2)
    add_para(doc, "Adaptive capacity shows both strengths and weaknesses. All households have bank accounts and access to a primary school and PHC, while many have ration cards and group membership. However, savings are very low, alternative livelihood options are weak and no household reported receiving mining compensation.")
    add_table(doc, ["Adaptive capacity indicator", "Finding"], [
        ["Savings", "5 households (8.3%)"],
        ["Bank account", "60 households (100.0%)"],
        ["SHG/union/group membership", "46 households (76.7%)"],
        ["Ration card/social protection", "40 households (66.7%)"],
        ["Mining compensation", "0 households (0.0%)"],
        ["Alternative livelihood", "No clear alternative livelihood for most households"],
        ["PHC access", "60 households (100.0%)"],
    ], [6, 9])
    add_heading(doc, "5.5 Community Perceptions and Open-ended Responses", 2)
    add_para(doc, "The open-ended responses make the survey findings more concrete. Respondents describe mining impacts through everyday experiences: water scarcity, unsafe water, skin problems, dust, noise, damaged roads, vibration during blasting, cracks in houses and lack of secure local employment.")
    add_table(doc, ["Major problem theme", "Relative mention count"], [[a, b] for a, b in problems[:6]], [10, 4])
    add_table(doc, ["Support requested by respondents", "Relative mention count"], [[a, b] for a, b in supports[:5]], [10, 4])

    add_heading(doc, "Chapter 6: Composite Vulnerability Index", 1)
    add_para(doc, "The Composite Vulnerability Index combines environmental exposure, sensitivity and adaptive capacity deficit into one household-level score. The overall vulnerability score ranges from 0.05 to 0.33, with an average of 0.19. This indicates that vulnerability in Mandamarri is moderate on average, but some households face much higher risk due to overlapping environmental, health and livelihood pressures.")
    add_picture(doc, chart1, "Figure 3: Composite vulnerability levels among households")
    add_table(doc, ["Vulnerability level", "Number of households", "Percentage"], [[k, v, f"{pct(v)}%"] for k, v in vuln_levels.items()], [5, 5, 4])
    add_heading(doc, "6.1 Interpretation", 2)
    add_para(doc, "The index results show that vulnerability is not limited to households working directly in mines. Non-mining households also experience dust, noise and water-quality problems. The most vulnerable households are those where environmental exposure combines with low savings, debt, health problems, unstable income, dependents and lack of compensation.")
    add_para(doc, "Exposure is high because dust, noise and water problems are nearly universal. Sensitivity varies more because households differ in education, health, land, income stability and debt. Adaptive capacity is mixed: bank accounts and public services exist, but lack of savings, lack of mining compensation and limited alternative livelihoods weaken resilience.")

    add_heading(doc, "Chapter 7: Major Findings and Recommendations", 1)
    add_heading(doc, "7.1 Major Findings", 2)
    add_bullets(doc, [
        "Mining is a major livelihood base in Mandamarri, but benefits are uneven across households.",
        "Dust pollution, regular noise disturbance and water-quality problems were reported by all surveyed households.",
        "A large share of households depend on one main earning member, increasing economic risk.",
        "Health insecurity is visible through illness affecting earning capacity, hospitalization and limited insurance coverage.",
        "No household reported receiving mining compensation, even though open-ended responses mention house cracks and infrastructure damage.",
        "Only 8.3 percent of households reported savings, which weakens crisis management capacity.",
        "The overall vulnerability index average is 0.19, with some households facing high vulnerability due to multiple overlapping pressures.",
    ])
    add_heading(doc, "7.2 Recommendations", 2)
    add_bullets(doc, [
        "Establish a regular safe drinking-water system, including filtration and summer water-scarcity planning.",
        "Conduct technical inspection of houses affected by blasting vibration and provide transparent compensation for verified damage.",
        "Give priority to local youth in mining-related employment, skill training and contract work.",
        "Strengthen dust-control measures through water sprinkling, road maintenance, covered transport and monitoring of open-cast dust.",
        "Improve local roads and transport services, especially where heavy mining vehicles damage roads.",
        "Organize periodic health camps for respiratory illness, skin disease, hearing problems, back pain and occupational risks.",
        "Expand social protection and insurance coverage for vulnerable households, especially those with illness, debt and dependents.",
        "Create a community grievance redress mechanism involving residents, mine authorities and local government.",
        "Support alternative livelihoods such as small enterprise training, livestock support, tailoring, driving, repair services and self-help group enterprises.",
    ])

    add_heading(doc, "Chapter 8: Conclusion", 1)
    add_para(doc, "The study shows that mining in Mandamarri has a dual character. It provides employment and supports the local economy, but it also creates environmental, health, housing and livelihood-related pressures. The most serious impacts reported by households are dust pollution, noise disturbance, water-quality problems, water scarcity, damaged roads, house cracks and lack of secure local employment.")
    add_para(doc, "The Composite Vulnerability Index confirms that vulnerability is multidimensional. Households are not vulnerable only because they live near mines; they become vulnerable when exposure combines with low savings, unstable work, illness, debt, dependents and weak compensation systems. Therefore, reducing mining impact requires more than employment generation. It requires safe water, pollution control, health protection, compensation, alternative livelihoods and meaningful community participation.")
    add_para(doc, "A people-centred mining policy for Mandamarri should ensure that local communities receive both economic opportunities and protection from environmental harm. Development will be meaningful only when the costs of mining are reduced for the households who live closest to it.")

    add_heading(doc, "References", 1)
    refs = [
        "Blaikie, P., Cannon, T., Davis, I., & Wisner, B. (1994). At Risk: Natural Hazards, People's Vulnerability and Disasters. Routledge.",
        "Chambers, R. (1989). Vulnerability, coping and policy. IDS Bulletin, 20(2), 1-7.",
        "International Labour Organization. (1999). Social and Labour Issues in Small-scale Mines. ILO.",
        "IPCC. (2014). Climate Change 2014: Impacts, Adaptation, and Vulnerability. Cambridge University Press.",
        "Ministry of Coal, Government of India. Annual reports and coal sector documents.",
        "Singareni Collieries Company Limited. Public information and company reports on coal mining operations.",
        "Primary household survey conducted in Mandamarri, Mancherial district, 2026.",
    ]
    add_bullets(doc, refs)

    add_heading(doc, "Appendix A: Survey Dimensions Used", 1)
    add_table(doc, ["Section", "Variables covered"], [
        ["Basic information", "Village, respondent, age, gender, religion, caste, house type, household size"],
        ["Education and health", "Education, vocational training, chronic illness, hospitalization, medical expense, health insurance"],
        ["Livelihood", "Main livelihood, active members, mining work, income source, income stability, employment type, documentation"],
        ["Exposure", "Distance from mine, dust, noise, accident, water quality, agriculture/livestock impact, relocation"],
        ["Sensitivity", "Land, livestock, water access, toilet, debt, loan source, social marginalization, dependents"],
        ["Adaptive capacity", "Savings, bank, group membership, social protection, compensation, alternative livelihood, PHC, school"],
        ["Perceptions/open-ended", "Impact severity, awareness, crisis confidence, decision-making, problems and support required"],
    ], [5, 10])

    add_heading(doc, "Appendix B: Data Source Note", 1)
    add_para(doc, "All tables and statistics in this final draft are derived from the Excel workbook named PavanThesis.0.xlsx. Some response spellings and category labels were standardized only for readability in the report, while the interpretation remains based on the original survey data.")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Impact of Mining on Local Communities - Mandamarri, Mancherial"

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
