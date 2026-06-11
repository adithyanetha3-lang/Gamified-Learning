from pathlib import Path
from collections import Counter
import json
import re

import openpyxl
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path("outputs/pavan_mandamarri")
SOURCE = BASE / "Impact_of_Mining_on_Local_Communities_Mandamarri_Mancherial_Final_Draft.docx"
ANALYSIS = BASE / "analysis.json"
XLSX = Path(r"C:/Users/adith/AppData/Local/Packages/5319275A.WhatsAppDesktop_cv1g1gvanyjgm/LocalState/sessions/F53961CCE067D31894D85901B81F73D5755EAA99/transfers/2026-23/PavanThesis.0.xlsx")
OUT = BASE / "Impact_of_Mining_on_Local_Communities_Mandamarri_Mancherial_Expanded_40_50_Page_Project.docx"


def clean(v):
    if v is None:
        return ""
    return re.sub(r"\s+", " ", str(v).strip())


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(7)
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color="FFFFFF")
        set_cell_shading(cell, "1F4E79")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def configure(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.1)
        section.left_margin = Cm(2.6)
        section.right_margin = Cm(2.1)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for style in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style].font.name = "Times New Roman"
        styles[style].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def append_source_doc(target, source):
    for element in source.element.body:
        if element.tag.endswith("sectPr"):
            continue
        target.element.body.append(element)


def rows_for(wb, sheet):
    ws = wb[sheet]
    data = list(ws.iter_rows(values_only=True))
    headers = [clean(x) or f"Column {i+1}" for i, x in enumerate(data[0])]
    rows = []
    for row in data[1:]:
        if any(clean(x) for x in row):
            rows.append({headers[i]: row[i] if i < len(row) else None for i in range(len(headers))})
    return headers, rows


def freq(rows, key, limit=20):
    c = Counter(clean(r.get(key)) for r in rows if clean(r.get(key)))
    total = sum(c.values()) or 1
    return [[k, v, f"{v * 100 / total:.1f}%"] for k, v in c.most_common(limit)]


def stats(rows, key):
    vals = []
    for r in rows:
        v = r.get(key)
        if isinstance(v, (int, float)):
            vals.append(float(v))
        elif clean(v):
            m = re.search(r"\d+(?:,\d{3})*(?:\.\d+)?", clean(v))
            if m:
                vals.append(float(m.group(0).replace(",", "")))
    if not vals:
        return []
    vals = sorted(vals)
    return [
        ["Count", len(vals)],
        ["Minimum", vals[0]],
        ["Maximum", vals[-1]],
        ["Mean", f"{sum(vals) / len(vals):.2f}"],
        ["Median", f"{vals[len(vals)//2] if len(vals)%2 else (vals[len(vals)//2-1]+vals[len(vals)//2])/2:.2f}"],
    ]


def add_long_section(doc, title, paragraphs):
    add_heading(doc, title, 2)
    for para in paragraphs:
        add_para(doc, para)


def main():
    data = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    wb = openpyxl.load_workbook(XLSX, data_only=True)
    doc = Document(SOURCE)
    configure(doc)
    doc.add_page_break()

    add_heading(doc, "Expanded Project Chapters and Detailed Discussion", 1)
    add_para(doc, "The following expanded chapters convert the compact final draft into a full-length academic project report. They provide deeper conceptual explanation, detailed interpretation of the field data, chapter-wise discussion, and appendices suitable for a 40-50 page submission.")

    add_heading(doc, "Expanded Chapter 1: Detailed Introduction", 1)
    add_long_section(doc, "1.1 Mining and Local Development", [
        "Coal mining is often presented as a driver of regional development because it creates jobs, supports transport networks, generates revenue and supplies energy for wider industrial growth. In coal-belt regions, mining also shapes the settlement pattern, labour market, public infrastructure and aspirations of local households. Mandamarri, located in Mancherial district of Telangana, reflects this close relationship between extractive industry and community life.",
        "However, the development contribution of mining cannot be understood only through employment. People living near mining areas also experience the everyday costs of extraction. These include dust pollution, blasting vibration, noise, water contamination, road damage, traffic risk, damage to houses, and pressure on local health and transport facilities. Such effects are not always captured in official production or employment statistics, but they are central to household well-being.",
        "In Mandamarri, households do not experience mining as a distant industrial activity. They experience it through water supply, road conditions, house structures, job opportunities, daily travel, health expenditure and uncertainty about compensation. Therefore, this project treats mining as a social, environmental and economic process that enters ordinary household life in multiple ways.",
        "A local community perspective is important because mining impacts are uneven. Some households may receive secure employment, while others receive only indirect or temporary income. Some houses may be closer to blasting or heavy vehicle movement than others. Households with savings and stable income may cope better, while households with illness, debt and dependents may become more vulnerable. This unevenness is the main reason for using a vulnerability framework in the study.",
    ])
    add_long_section(doc, "1.2 Need for the Study", [
        "The need for this study emerges from the gap between mining-led economic growth and community-level lived experience. In many mining regions, local people recognize that mining provides employment, but they also report that basic living conditions are affected. When dust, noise and water problems become regular, households begin to carry the hidden costs of mining through medical expenses, repair costs, travel difficulties and reduced quality of life.",
        "Mandamarri households in the survey repeatedly mentioned water scarcity, unfiltered water, dust pollution, road damage, blasting vibration, cracks in houses and shortage of local job opportunities. These problems show that mining impact is not limited to one sector. It affects public health, settlement infrastructure, livelihood security and local governance.",
        "A systematic household-level study is useful because it converts scattered experiences into evidence. The survey of 60 households gives a clearer picture of how common specific problems are, which groups are more exposed, and what forms of support residents expect. The study also helps identify practical interventions instead of treating mining impact as a general complaint.",
    ])
    add_long_section(doc, "1.3 Key Concepts Used in the Study", [
        "Impact refers to the positive and negative changes created by mining in the life of local communities. In this report, impact includes livelihood opportunities, income sources, environmental pollution, health risk, infrastructure problems, asset damage, social protection and community perceptions.",
        "Exposure refers to the degree to which households come into contact with mining-related hazards. Indicators such as dust pollution, noise disturbance, distance from mine, accident experience and water-quality issues represent exposure.",
        "Sensitivity refers to the household conditions that make the effect of exposure stronger. Low education, illness, debt, dependents, landlessness and unstable income can increase sensitivity because these factors reduce the ability of a household to absorb shock.",
        "Adaptive capacity refers to the resources and support systems that help households manage stress. Savings, bank accounts, self-help group membership, ration cards, compensation, health facilities, schools and alternative livelihood options are important adaptive-capacity indicators.",
    ])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 2: Review of Literature", 1)
    literature = [
        ("2.1 Mining and the Development Paradox", [
            "The literature on mining frequently describes a development paradox. Mining regions can contribute significantly to economic output, but communities near extraction sites may face pollution, occupational hazards and social insecurity. The paradox occurs when benefits are spread across the economy while costs are concentrated locally.",
            "In coal mining regions, this paradox is visible in the relationship between employment and environmental harm. A household may depend on mining wages and still suffer from dust, noise, unsafe water or house damage. This means that mining cannot be evaluated only by counting jobs; it must also be evaluated by examining the quality of life of those who live near mines.",
        ]),
        ("2.2 Environmental Effects of Coal Mining", [
            "Coal mining generates dust through excavation, transport, dumping, blasting and movement of heavy vehicles. Dust affects air quality and may contribute to respiratory and skin-related problems. Noise from machinery, blasting and vehicles can disturb daily life, sleep and mental well-being.",
            "Water-related impacts are also widely discussed in mining studies. Mining can affect groundwater levels, surface drainage and water quality. Even where water access exists through taps or borewells, residents may still perceive the water as unsafe if it is contaminated, irregular or unfiltered.",
        ]),
        ("2.3 Livelihood and Labour Issues", [
            "Mining creates both direct and indirect employment. Direct employment includes permanent and contract mine work, while indirect dependence includes drivers, vendors, mechanics, small shops, transport services and daily wage labour. The stability and security of these livelihoods differ significantly.",
            "Contract and informal workers are often more vulnerable because their income may be irregular, their benefits weaker and their bargaining power limited. Non-mining households may also be affected because agriculture, livestock and small businesses can suffer from pollution, road damage or reduced local demand.",
        ]),
        ("2.4 Health and Social Protection", [
            "Health vulnerability is a major concern in mining areas. Respiratory illness, skin disease, back pain, accident risk and occupational injuries may increase household expenditure and reduce earning capacity. When a main earning member becomes sick, the effect spreads to the entire household.",
            "Social protection can reduce vulnerability, but only when it is accessible and sufficient. Ration cards, pensions, health insurance, compensation and public health services are important. However, if compensation for mining-related damage is absent, households may carry private costs for a public industrial activity.",
        ]),
        ("2.5 Vulnerability Framework", [
            "The vulnerability framework is useful because it explains why all households do not suffer equally. Exposure may be common across a settlement, but sensitivity and adaptive capacity vary. A household with savings, secure employment and lower debt can recover more easily than a household with illness, dependents and unstable income.",
            "This project uses exposure, sensitivity and adaptive capacity as the conceptual foundation for the Composite Vulnerability Index. The index is not a replacement for qualitative experience; rather, it organizes household-level risk into measurable components that can support discussion and planning.",
        ]),
    ]
    for heading, paras in literature:
        add_long_section(doc, heading, paras)
    add_table(doc, ["Author/Framework", "Relevance to present study", "Application in Mandamarri"], [
        ["Vulnerability approach", "Explains risk as a combination of exposure, sensitivity and coping ability", "Used to classify household-level mining vulnerability"],
        ["Political ecology perspective", "Shows how environmental costs and economic benefits are distributed unevenly", "Useful for understanding why local households report costs despite mining employment"],
        ["Livelihood framework", "Examines assets, income sources and coping strategies", "Used to analyse mining, agriculture, daily wage work and small business dependence"],
        ["Environmental justice approach", "Focuses on who bears pollution and who receives benefits", "Relevant because households report dust, noise, water and house damage concerns"],
    ], [4, 6, 6])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 3: Detailed Methodology", 1)
    method_sections = [
        ("3.1 Research Design", [
            "The research design is descriptive, analytical and field-data based. The descriptive part records the socio-economic profile of households, including age, gender, caste, religion, household size, education, livelihood and housing. The analytical part examines relationships among mining exposure, health, debt, income stability and adaptive capacity.",
            "The design is suitable for a local project because it allows both numerical summary and interpretation of lived experience. Frequencies and percentages show the scale of problems, while open-ended responses explain what those problems mean in daily life.",
        ]),
        ("3.2 Sampling and Respondents", [
            "The study uses a household survey of 60 respondents from Mandamarri. The household was selected as the unit of analysis because mining impacts are experienced collectively through income, health expenditure, water use, housing, transport and coping strategies.",
            "The sample includes households directly connected with mining and households involved in non-mining livelihoods. This makes it possible to understand mining as a community-level influence rather than as a workplace-only issue.",
        ]),
        ("3.3 Tools of Data Collection", [
            "The survey schedule was divided into sections covering basic information, education and health, livelihood, environmental exposure, socio-economic sensitivity, adaptive capacity, shocks, perceptions, assets and open-ended questions.",
            "Open-ended questions were included to capture issues that may not fit neatly into predefined categories. These responses were especially important for identifying repeated concerns such as unsafe water, damaged roads, blasting vibration, house cracks, dust and lack of local employment.",
        ]),
        ("3.4 Data Processing", [
            "After data collection, responses were entered into the Excel workbook. Category frequencies, percentages, scores and index values were calculated. Some spelling variations were standardized for report readability, but the interpretation remains based on the original survey responses.",
            "For the Composite Vulnerability Index, indicators were organized under exposure, sensitivity and adaptive capacity. Higher scores indicate greater vulnerability or lower coping ability depending on the indicator. The final vulnerability level was then classified into low, medium and high categories.",
        ]),
        ("3.5 Limitations", [
            "The study is limited to 60 households and therefore does not represent every household in Mandamarri or all mining regions of Telangana. It is a local-level academic study based on available primary data.",
            "The study relies partly on self-reported experiences. Perceptions of pollution, water quality and damage are important because they shape household behaviour, but technical environmental testing was not conducted within the scope of this project.",
        ]),
    ]
    for heading, paras in method_sections:
        add_long_section(doc, heading, paras)
    add_table(doc, ["Survey section", "Purpose", "Examples of variables"], [
        ["Basic information", "To understand household profile", "Age, gender, caste, religion, household size, house type"],
        ["Education and health", "To assess human capital and health risk", "Education, vocational training, chronic illness, hospitalization, health insurance"],
        ["Livelihood", "To study economic dependence and income security", "Main work, mining employment, income source, income stability"],
        ["Exposure", "To measure direct mining-related pressures", "Dust, noise, distance from mine, accident, water quality"],
        ["Sensitivity", "To identify conditions that increase vulnerability", "Land, debt, dependents, livestock, water, toilet"],
        ["Adaptive capacity", "To measure coping resources", "Savings, bank, group membership, compensation, health facility"],
        ["Perceptions", "To understand subjective severity and governance", "Impact severity, rights awareness, crisis confidence, decision-making"],
    ], [4, 5.5, 6])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 4: Detailed Data Tables", 1)
    basic_h, basic = rows_for(wb, "Basic Information")
    edu_h, edu = rows_for(wb, "Section_B_Education_Health")
    liv_h, liv = rows_for(wb, "Section_C_Livelihood")
    exp_h, exp = rows_for(wb, "Section_D_Exposure")
    sens_h, sens = rows_for(wb, " Sensitivity (socio-economic v")
    adap_h, adap = rows_for(wb, "Adaptive Capacity")
    shock_h, shock = rows_for(wb, "Shocks, losses & coping strateg")
    percep_h, percep = rows_for(wb, "Perceptions, awareness & partic")
    asset_h, asset = rows_for(wb, "Assets & infrastructure")
    sheet10_h, sheet10 = rows_for(wb, "Sheet10")

    table_specs = [
        ("4.1 Caste Composition", ["Caste", "Households", "Percentage"], freq(basic, "Caste")),
        ("4.2 Housing Type", ["House Type", "Households", "Percentage"], freq(basic, "HH_type")),
        ("4.3 Education Level", ["Education", "Households", "Percentage"], freq(edu, "B1_Education")),
        ("4.4 Main Livelihood", ["Livelihood", "Households", "Percentage"], freq(liv, "C1_Main_Livelihood")),
        ("4.5 Main Income Source", ["Income Source", "Households", "Percentage"], freq(liv, "C5_Main sources of income")),
        ("4.6 Income Stability", ["Income Stability", "Households", "Percentage"], freq(liv, "C7_Income Stability")),
        ("4.7 Distance from Mine", ["Distance", "Households", "Percentage"], freq(exp, "D1_Distance_Mine")),
        ("4.8 Land Ownership", ["Land Status", "Households", "Percentage"], freq(sens, "E1_ Land ownership")),
        ("4.9 Loan Source", ["Loan Source", "Households", "Percentage"], freq(sens, "E5b_ Main source of loan")),
        ("4.10 Social Protection Access", ["Protection", "Households", "Percentage"], freq(adap, "F4_Access to government social protection")),
        ("4.11 Recent Major Shock", ["Shock", "Households", "Percentage"], freq(shock, "G2_Most recent major shock")),
        ("4.12 Perceived Impact Severity", ["Severity", "Households", "Percentage"], freq(percep, "H1_Impact_Severity")),
        ("4.13 Confidence to Manage Crisis", ["Confidence", "Households", "Percentage"], freq(percep, "H4_ How confident are you that the household can manage a major crisis without external help?")),
        ("4.14 Transport Ownership", ["Transport", "Households", "Percentage"], freq(asset, "I4. Means of transport owned")),
        ("4.15 Vulnerability Level", ["Level", "Households", "Percentage"], freq(sheet10, "Vulnerability Level")),
    ]
    for title, headers, rows in table_specs:
        add_heading(doc, title, 2)
        add_table(doc, headers, rows, [6, 4, 4])
        add_para(doc, f"The table on {title.lower()} shows the distribution of households across important categories. These categories help identify the social and economic background through which mining impacts are experienced in Mandamarri.")

    add_heading(doc, "4.16 Descriptive Statistics", 2)
    add_table(doc, ["Variable", "Statistic", "Value"], [["Age", a, b] for a, b in stats(basic, "Age")] + [["Household Size", a, b] for a, b in stats(basic, "HH Size")] + [["Income", a, b] for a, b in stats(liv, "C6_Income Category")], [5, 5, 4])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 5: Detailed Interpretation and Discussion", 1)
    discussion = [
        ("5.1 Mining as Employment and Risk", [
            "The survey shows that mining is a major part of Mandamarri's livelihood structure. A majority of households are currently connected to mining, either directly or indirectly. This confirms that mining cannot be treated only as an external industry; it is tied to household income, local aspiration and community identity.",
            "At the same time, dependence on mining creates risk. If jobs are contract-based, informal or dependent on one earning member, households may face insecurity despite being connected to the mining economy. This is why the report distinguishes between employment access and livelihood security. Employment exists, but it does not automatically remove vulnerability.",
            "Households outside mining also face mining-related impacts. Farmers, drivers, vendors and daily wage labourers may not receive direct mining benefits, but they live with dust, noise, traffic and water problems. This makes mining impact a community issue rather than only a worker issue.",
        ]),
        ("5.2 Water, Dust and Noise as Everyday Problems", [
            "The most striking finding is the universality of environmental complaints. Dust pollution, regular noise and water-quality problems were reported across the survey. This suggests that environmental exposure is not an isolated problem affecting only a few houses; it is a shared condition in the study area.",
            "Open-ended responses show that water is both an environmental and public health issue. Respondents described unfiltered water, water scarcity in summer and skin-related problems. Even when tap connections are present, the quality and reliability of water remain central concerns.",
            "Dust and noise affect quality of life in less visible but continuous ways. Dust may settle on houses, roads, vegetation and water sources, while regular noise and blasting vibration disturb the settlement environment. These impacts can reduce comfort, increase stress and create fear about long-term health.",
        ]),
        ("5.3 Housing and Infrastructure Damage", [
            "Several respondents mentioned cracks in houses and vibrations due to blasting. Housing damage is an important issue because it converts environmental exposure into direct financial loss. Poor and semi-pucca households may find it difficult to repair damage without compensation.",
            "Damaged roads and heavy vehicle movement were also raised in open-ended responses. Roads are not only transport infrastructure; they affect access to schools, health facilities, work, markets and emergency services. When mining vehicles damage roads, the wider community carries the cost through travel difficulty and safety risk.",
        ]),
        ("5.4 Health Vulnerability", [
            "Health vulnerability appears through chronic illness, hospitalization and medical expenses. When illness affects earning capacity, the household faces a double burden: income decreases while expenditure increases. This burden is more severe where savings are low.",
            "Only a section of households reported health insurance. The absence of universal health protection increases household sensitivity to mining-related and non-mining-related health shocks. Periodic health camps and stronger PHC referral systems would reduce this risk.",
        ]),
        ("5.5 Debt, Savings and Coping", [
            "Debt was reported by a significant share of households, while savings were reported by only a small minority. This combination is important because it reveals weak financial resilience. When households face illness, repair costs or income loss, they may need to borrow from banks, relatives or moneylenders.",
            "The coping strategy data show that serious illness and economic shocks can lead to loans or asset sale. Such coping strategies may solve immediate problems but can increase long-term vulnerability. Therefore, financial resilience should be treated as an important part of mining-area development.",
        ]),
        ("5.6 Governance and Participation", [
            "The absence of reported mining compensation is one of the most important findings. When households report house cracks, water problems or infrastructure damage but do not receive compensation, trust in mining governance can weaken.",
            "Community participation is necessary for fair decision-making. Local residents should be informed about rights, environmental standards, compensation procedures and grievance redress mechanisms. Without participation, mining-affected communities may remain dependent on informal complaints rather than formal solutions.",
        ]),
    ]
    for heading, paras in discussion:
        add_long_section(doc, heading, paras)
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 6: Composite Vulnerability Index - Detailed Explanation", 1)
    cvi_paras = [
        "The Composite Vulnerability Index is used in this project to summarize the household-level risk pattern. It combines three dimensions: Exposure Index, Sensitivity Index and Adaptive Capacity Deficit Index. Each dimension captures a different aspect of vulnerability.",
        "Exposure Index measures direct contact with mining-related hazards. In Mandamarri, exposure is high because dust, noise and water-quality problems were reported by all households. Accident experience and agriculture/livestock impact vary, but the general environmental pressure is widespread.",
        "Sensitivity Index measures the conditions that make households more likely to be affected. A household with dependents, debt, illness, unstable income and no land is more sensitive than a household with stable income, fewer dependents and stronger assets.",
        "Adaptive Capacity Deficit Index measures the absence or weakness of coping resources. In Mandamarri, bank accounts and public services exist, but lack of savings, absence of compensation and limited alternative livelihood reduce the ability of households to recover from shocks.",
        "The overall vulnerability index average is about 0.19. The minimum value is 0.05 and the maximum value is 0.33. This range shows that vulnerability is moderate overall, but it is not equal for every household.",
    ]
    for para in cvi_paras:
        add_para(doc, para)
    add_table(doc, ["Component", "Role in vulnerability", "Mandamarri interpretation"], [
        ["Exposure", "Shows direct pressure from mining hazards", "Dust, noise and water problems are nearly universal"],
        ["Sensitivity", "Shows household weakness under pressure", "Debt, dependents, illness and unstable income increase risk"],
        ["Adaptive capacity deficit", "Shows weak coping and recovery resources", "Low savings and no compensation are major concerns"],
        ["Overall CVI", "Combines the three dimensions", "Moderate average vulnerability with some high-risk households"],
    ], [4.5, 5, 6])
    add_heading(doc, "6.1 Vulnerability-Level Interpretation", 2)
    add_para(doc, "Low vulnerability households generally have lower sensitivity and stronger coping conditions. Medium vulnerability households experience significant exposure but may have some resources such as stable income, bank access or social protection. High vulnerability households face overlapping risk: environmental exposure, weak savings, health problems, debt, dependents, unstable income or lack of compensation.")
    add_para(doc, "The vulnerability-level classification helps policy planning because it identifies which households need urgent support. High vulnerability households should receive priority in health screening, compensation verification, social protection linkage and alternative livelihood support.")
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 7: Policy Suggestions and Action Plan", 1)
    add_table(doc, ["Problem", "Evidence from study", "Recommended action", "Responsible actors"], [
        ["Water quality and scarcity", "100% reported water-quality problems; open-ended responses mention unfiltered water and summer scarcity", "Install filtration, ensure regular water supply, test water quality periodically", "Gram Panchayat, SCCL, district administration"],
        ["Dust pollution", "100% reported dust or airborne pollution", "Water sprinkling, covered transport, road cleaning, green buffer, dust monitoring", "Mine authorities, pollution control board"],
        ["Noise and blasting vibration", "100% reported regular noise; responses mention cracks and vibration", "Blasting schedule communication, vibration monitoring, technical house inspection", "Mine authorities, revenue department"],
        ["House damage", "Open-ended responses mention cracks in houses", "Transparent damage survey and compensation mechanism", "District administration, mine company"],
        ["Limited local employment", "Respondents request local job priority and reservation", "Skill training, local contract preference, apprenticeship, placement support", "SCCL, labour department, skill mission"],
        ["Weak savings and debt", "Only 8.3% reported savings; 35% reported debt", "Financial literacy, SHG strengthening, emergency credit at low interest", "Banks, SERP, SHGs"],
        ["Health risks", "Hospitalization and illness affecting earning reported", "Regular health camps, insurance enrolment, referral support", "Health department, PHC, mine medical wing"],
        ["Transport and road damage", "Respondents mention damaged roads and weak transport", "Road repair, heavy vehicle route management, regular bus/auto connectivity", "R&B department, local bodies"],
    ], [3.5, 5, 5, 3.5])
    add_para(doc, "The action plan should be implemented through a participatory mechanism. Community representatives, local elected members, women group members, youth, mine officials and district administration should meet periodically to review grievances and progress. This will improve transparency and reduce conflict.")
    add_heading(doc, "7.1 Priority Interventions", 2)
    add_numbered(doc, [
        "Prepare a household-wise list of houses affected by blasting vibration and cracks.",
        "Conduct water-quality testing in all main drinking-water sources in Mandamarri.",
        "Start monthly dust and road-condition monitoring near mining transport routes.",
        "Organize health screening camps focusing on respiratory, skin, hearing and musculoskeletal issues.",
        "Create local youth skill training linked to mining and non-mining jobs.",
        "Strengthen social protection coverage for households with illness, dependents, debt and no savings.",
        "Create a formal grievance register and publicly display action taken on complaints.",
    ])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 8: Chapter-wise Synthesis", 1)
    add_long_section(doc, "8.1 Synthesis of Socio-Economic Profile", [
        "The socio-economic profile of the respondents shows that Mandamarri is not a uniform settlement. Households differ in caste background, education, livelihood, housing type, debt status and coping resources. This variation matters because mining impact is filtered through household conditions. A household with better education, secure income and savings can manage stress more easily than a household with illness, debt and one earning member.",
        "The dominance of semi-pucca houses in the sample is important in relation to blasting vibration and house cracks. Semi-pucca houses may be more vulnerable to structural damage than stronger pucca buildings. When respondents report cracks and damage, the issue should be examined technically rather than dismissed as a general complaint.",
        "The age and household-size profile also indicates that many households are in active working age but have dependents. Dependents increase responsibility for food, education, health care and daily expenses. When mining-related problems increase household expenditure, the pressure is felt across the family.",
        "Education levels show that many respondents have completed higher secondary education, and some have graduation or technical training. This creates an opportunity for skill-based employment and local placement. However, education alone does not guarantee livelihood security when local job access is limited or contract-based.",
    ])
    add_long_section(doc, "8.2 Synthesis of Livelihood Findings", [
        "Mining is a central livelihood source in the study area, but the survey shows that Mandamarri's economy is mixed. Households are connected to mining, agriculture, small business, driving, daily wage work, security work and private jobs. This mixed pattern shows that mining has direct and indirect influence on the local economy.",
        "The key livelihood concern is not only whether people have work, but whether work is stable, secure and sufficient. Households depending on contract work, seasonal activity or one earning member remain vulnerable. Any health shock or job interruption can quickly affect the household budget.",
        "The data also shows that formal documentation exists for many households. This is a positive sign because documentation can help households access welfare schemes, employment records and financial services. But documentation must be linked to actual support, especially in compensation and social protection.",
        "A sustainable livelihood strategy for Mandamarri should not depend only on mining jobs. It should include mining-linked skill training, non-mining enterprise support, transport services, repair work, tailoring, livestock support, small businesses and women's self-help group enterprises. Diversification will reduce the risk of dependence on a single industry.",
    ])
    add_long_section(doc, "8.3 Synthesis of Environmental Findings", [
        "Environmental exposure is the strongest evidence in the study. Dust, regular noise and water-quality problems were reported by all surveyed households. This means that mining impact is not experienced by a small section alone. The settlement environment as a whole is affected.",
        "Dust pollution is linked with both mining operations and transport movement. Dust can enter houses, settle on roads, affect children and elderly people, and create a perception of continuous environmental decline. Even if households depend on mining employment, they may still demand stronger dust-control measures.",
        "Water-quality concern is particularly serious because safe water is a basic requirement for health and dignity. Respondents referred to unfiltered water, water scarcity and skin-related problems. These issues require scientific water testing, filtration support and regular public reporting of water quality.",
        "Noise and blasting vibration affect the sense of safety in the settlement. When people report cracks in houses, they are not only reporting physical damage; they are expressing anxiety about living near an industrial hazard. Monitoring, communication and compensation procedures are necessary to reduce this fear.",
    ])
    add_long_section(doc, "8.4 Synthesis of Health and Coping Findings", [
        "Health problems reduce household resilience. The survey records illness affecting earning, hospitalization and medical expenditure. A mining-affected household with illness faces multiple burdens: reduced labour capacity, medical costs, travel for treatment and possible debt.",
        "Health insurance coverage is not universal, which means many households may depend on savings or loans during illness. Since savings are very low in the study area, health shocks can push households into debt. This makes health protection an economic issue as well as a medical issue.",
        "The coping data suggests that households may take loans or sell assets during serious shocks. These strategies can help immediately but weaken long-term security. When assets are sold, future earning capacity may decline. When loans are taken, monthly repayment pressure increases.",
        "Regular health camps, occupational-risk screening, insurance enrolment and referral support should become part of mining-area community development. Health interventions should focus on respiratory problems, skin disease, hearing, back pain, accident risk and chronic illness.",
    ])
    add_long_section(doc, "8.5 Synthesis of Governance Findings", [
        "Governance is visible in the gap between reported problems and received support. No household reported mining compensation, even though open-ended responses mention house cracks, water issues and damage. This gap can create dissatisfaction and mistrust.",
        "A transparent grievance system is necessary. Residents should know where to report mining-related damage, how verification is done, what documents are required, and how compensation decisions are made. A grievance register should record complaints and action taken.",
        "Community participation should include women, youth, elderly people, workers, non-mining households and marginalized groups. Mining impact is not only a worker issue, so participation should not be limited to mine employees.",
        "Local governance should connect mining authorities with Gram Panchayat, district administration, health department, roads department, pollution control authorities and community organizations. A coordinated approach is needed because mining impact crosses departmental boundaries.",
    ])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 9: Thematic Case Narratives from Field Responses", 1)
    narratives = [
        ("9.1 Water as a Household Stress", [
            "Many open-ended responses describe water as one of the most urgent problems. Respondents referred to unfiltered water, water scarcity during summer and the need for regular water supply. This shows that water is not only an environmental issue but also a daily household-management issue.",
            "When drinking water is uncertain, households spend more time and effort securing water. Women, elderly persons and children often bear the everyday burden of water management. Water scarcity can also affect hygiene, cooking, livestock care and health.",
            "The demand for a filter bed and access to better water sources indicates that residents are not only complaining about mining but also proposing practical solutions. The policy response should therefore include water testing, filtration, storage and maintenance.",
        ]),
        ("9.2 House Cracks and Blasting Vibration", [
            "Several responses mention cracks in houses and vibrations due to blasting. For households, a house is both a shelter and a major asset. Damage to the house therefore affects security, dignity and financial stability.",
            "Blasting-related fear may continue even when visible damage is limited. Residents may worry about future damage, repair costs and safety of children and elderly persons. A technical survey by qualified engineers would help separate verified damage from general fear and create a fair basis for compensation.",
            "Compensation should be transparent and timely. If households believe that damage is ignored, conflict between communities and mining authorities may increase. A clear procedure can reduce mistrust.",
        ]),
        ("9.3 Roads, Heavy Vehicles and Mobility", [
            "Respondents mentioned damaged roads and weak transport facilities. Mining transport routes often carry heavy vehicles, which can damage roads and create dust. Poor roads affect school travel, health access, market movement and daily work.",
            "Transport is particularly important for households without private vehicles. If regular bus service is weak, residents depend on autos or informal transport. This increases cost and reduces mobility for poorer households.",
            "Road maintenance should be treated as part of mining-area responsibility. Where industrial traffic contributes to road damage, repair and dust control should be planned regularly.",
        ]),
        ("9.4 Employment Expectations and Local Youth", [
            "Open-ended responses repeatedly ask for local job priority, employment reservation and mining-related opportunities. These demands show that local residents expect mining to provide direct benefits to the community.",
            "Employment expectations are understandable because local communities live with the environmental costs of mining. If outsiders or only a small section receive jobs, resentment may increase. Local skill training and transparent recruitment can help address this concern.",
            "Employment policy should include both mining and non-mining opportunities. Technical training, driving, machine operation, repair services, safety training, small enterprise support and women-focused livelihood activities can reduce vulnerability.",
        ]),
        ("9.5 Health, Skin Disease and Pollution Anxiety", [
            "Some responses mention allergy, skin disease and health problems connected with water or pollution. Whether every illness is directly caused by mining requires medical assessment, but the perception itself is important because it shapes community trust and behaviour.",
            "A community health programme can address both diagnosis and awareness. Regular camps can identify common symptoms, refer serious cases and provide advice on water safety, dust exposure and occupational precautions.",
            "Health data should be collected over time. If respiratory, skin or hearing issues are common, targeted interventions can be designed. A one-time camp is useful, but periodic monitoring is better for mining areas.",
        ]),
    ]
    for heading, paras in narratives:
        add_long_section(doc, heading, paras)
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 10: Implementation Framework", 1)
    add_para(doc, "The recommendations of this project can be translated into a practical implementation framework. The framework below organizes actions into immediate, short-term and medium-term priorities. This makes the project more useful for planning because every recommendation is connected with time and responsible institutions.")
    add_table(doc, ["Time frame", "Priority action", "Expected result"], [
        ["Immediate", "Water-quality testing and drinking-water supply review", "Identification of unsafe sources and urgent corrective steps"],
        ["Immediate", "House-crack and blasting-vibration complaint registration", "Documented evidence for technical inspection"],
        ["Immediate", "Dust control on mining transport roads", "Reduced visible dust exposure"],
        ["Short term", "Health camps and insurance enrolment drive", "Reduced out-of-pocket health risk"],
        ["Short term", "Road repair and heavy vehicle route management", "Improved mobility and lower accident/dust risk"],
        ["Short term", "Local skill training for youth", "Better access to mining and non-mining employment"],
        ["Medium term", "Community grievance redress committee", "Improved trust and accountability"],
        ["Medium term", "Alternative livelihood and SHG enterprise support", "Reduced dependence on single industry"],
        ["Medium term", "Annual vulnerability monitoring", "Evidence-based planning for mining-affected households"],
    ], [3.5, 7, 6])
    add_long_section(doc, "10.1 Monitoring Indicators", [
        "Monitoring should not be limited to production or employment figures. A mining-affected settlement should also monitor safe water access, dust levels, road condition, house-damage complaints, health camp attendance, compensation claims, employment of local youth and grievance resolution.",
        "Household-level vulnerability monitoring can be repeated every year. The same exposure, sensitivity and adaptive capacity indicators may be used to see whether conditions improve or worsen. This will help shift mining governance from complaint response to preventive planning.",
    ])
    add_table(doc, ["Monitoring indicator", "Suggested frequency", "Why it matters"], [
        ["Water-quality test results", "Quarterly", "Shows whether drinking water is safe"],
        ["Dust-control compliance", "Monthly", "Reduces everyday exposure"],
        ["House-damage complaints", "Monthly", "Tracks blasting-related concerns"],
        ["Health camp records", "Quarterly", "Identifies common health risks"],
        ["Local employment numbers", "Half-yearly", "Tracks benefit sharing"],
        ["Compensation cases resolved", "Quarterly", "Improves accountability"],
        ["Road-condition inspection", "Monthly", "Supports mobility and safety"],
    ], [5, 4, 7])
    doc.add_page_break()

    add_heading(doc, "Expanded Chapter 11: Extended Conclusion", 1)
    for para in [
        "This project concludes that mining in Mandamarri is deeply connected with both livelihood opportunity and social vulnerability. The same mining economy that supports employment also produces environmental and infrastructural pressures that affect household well-being.",
        "The strongest evidence in the survey relates to environmental exposure. Dust pollution, regular noise disturbance and water-quality problems were reported across the sample. These are not minor inconveniences; they shape the everyday environment in which people cook, sleep, travel, work, study and manage health.",
        "Livelihood findings show that mining is important, but economic security is uneven. A household may be connected to mining and still remain vulnerable if employment is contractual, savings are low, health problems exist or dependents are many. Therefore, livelihood policy should focus on security, skills and diversification, not only job numbers.",
        "The Composite Vulnerability Index confirms the multidimensional character of the problem. Vulnerability emerges when exposure, sensitivity and weak adaptive capacity come together. Households with debt, illness, dependents, no savings and no compensation require special attention.",
        "A fair development approach for Mandamarri must reduce the costs carried by local communities. Safe water, pollution control, house-damage compensation, health protection, road repair, local employment priority and community participation should become central to mining governance.",
    ]:
        add_para(doc, para)
    doc.add_page_break()

    add_heading(doc, "Appendix C: Survey Schedule", 1)
    schedule = [
        ("Basic Information", basic_h),
        ("Education and Health", edu_h),
        ("Livelihood", liv_h),
        ("Environmental Exposure", exp_h),
        ("Sensitivity", list(rows_for(wb, " Sensitivity (socio-economic v")[0])),
        ("Adaptive Capacity", adap_h),
        ("Shocks, Losses and Coping", shock_h),
        ("Perceptions and Participation", percep_h),
        ("Assets and Infrastructure", asset_h),
    ]
    for title, headers in schedule:
        add_heading(doc, title, 2)
        add_numbered(doc, headers)

    add_heading(doc, "Appendix D: Sample Household-Level Vulnerability Data", 1)
    sample_rows = []
    for i, row in enumerate(sheet10[:30], start=1):
        sample_rows.append([
            i,
            clean(row.get("C1_Main_Livelihood")),
            f"{float(row.get('EI') or 0):.3f}",
            f"{float(row.get('SI') or 0):.3f}",
            f"{float(row.get('ADI') or 0):.3f}",
            f"{float(row.get('Vulnerability index') or 0):.3f}",
            clean(row.get("Vulnerability Level")),
        ])
    add_table(doc, ["S.No.", "Livelihood", "EI", "SI", "ADI", "VI", "Level"], sample_rows, [1.3, 5, 1.7, 1.7, 1.7, 1.7, 2.3])

    add_heading(doc, "Appendix E: Open-Ended Response Extracts", 1)
    open_h, open_rows = rows_for(wb, "Open-ended")
    extracts = []
    for i, row in enumerate(open_rows[:20], start=1):
        extracts.append([
            i,
            clean(row.get("J1. What do you think are the main problems your household faces because of mining?"))[:300],
            clean(row.get("J2. What support or change would help your household the most in reducing these problems?"))[:300],
        ])
    add_table(doc, ["S.No.", "Problems reported", "Support expected"], extracts, [1.2, 7, 7])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "Impact of Mining on Local Communities - Mandamarri, Mancherial"

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
